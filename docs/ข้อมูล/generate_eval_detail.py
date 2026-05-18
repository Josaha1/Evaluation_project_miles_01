#!/usr/bin/env python3
"""Generate detailed evaluation forms analysis document (.docx) with all questions"""

import json, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

JSON_PATH = os.path.join(os.environ.get('TEMP', '/tmp'), 'eval_full_data.json')

with open(JSON_PATH, 'r', encoding='utf-8') as f:
    evals = json.load(f)

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

style = doc.styles['Normal']
style.font.name = 'TH Sarabun New'
style.font.size = Pt(14)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')

def add_colored_heading(text, level=1, color=RGBColor(0, 51, 102)):
    h = doc.add_heading('', level=level)
    r = h.add_run(text)
    r.font.color.rgb = color
    return h

def set_cell(cell, text, bold=False, size=Pt(12), align=None):
    cell.text = ''
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.font.size = size
    r.font.bold = bold
    if align:
        p.alignment = align

# ══════════════════════════════════════════════
# Title
# ══════════════════════════════════════════════
title = doc.add_heading('', level=0)
run = title.add_run('รายงานวิเคราะห์แบบประเมิน 360 องศา (ฉบับละเอียด)')
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0, 51, 102)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('การนิคมอุตสาหกรรมแห่งประเทศไทย (กนอ.)\nปีงบประมาณ 2568 - 2569')
r.font.size = Pt(16)
r.font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph()

# ══════════════════════════════════════════════
# Section 1: Summary Table
# ══════════════════════════════════════════════
add_colored_heading('1. ภาพรวมแบบประเมินทั้งหมด')

tbl = doc.add_table(rows=1, cols=6)
tbl.style = 'Light Grid Accent 1'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['ID', 'ชื่อแบบประเมิน', 'ระดับ', 'ประเภท', 'คำถาม', 'ปีงบ']):
    set_cell(tbl.rows[0].cells[i], h, bold=True, size=Pt(13), align=WD_ALIGN_PARAGRAPH.CENTER)

fy_map = {1:'68', 2:'68', 3:'68', 4:'68', 5:'68', 14:'68',
          33:'69', 34:'69', 35:'69', 36:'69', 37:'69', 38:'69', 39:'69', 40:'69'}
type_map = {'internal':'ภายใน', 'external':'ภายนอก'}

for e in evals:
    eid = e['id']
    total_q = sum(len(q) for p in e['parts'] for a in p['aspects'] for q in [a['questions']])
    etype = type_map.get(e['type'], e['type'])
    if 'ตนเอง' in e['title']:
        etype = 'ตนเอง'
    short_title = e['title'].replace('แบบประเมิน 360 องศา ', '').replace('สำหรับ', '')
    row = tbl.add_row().cells
    for i, val in enumerate([str(eid), short_title[:40], e['grade'], etype, str(total_q), fy_map.get(eid, '?')]):
        set_cell(row[i], val, size=Pt(11), align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()

# ══════════════════════════════════════════════
# Section 2: Selection Logic
# ══════════════════════════════════════════════
add_colored_heading('2. วิธีการเลือกแบบประเมิน')

add_colored_heading('2.1 หลักการเลือก', level=2, color=RGBColor(0, 80, 120))
doc.add_paragraph('ระบบใช้ EvaluationLookupService เป็นศูนย์กลาง โดยพิจารณาจาก:')
for f in [
    'ระดับตำแหน่ง (Grade) ของผู้ถูกประเมิน — ตรวจสอบจาก grade_min/grade_max',
    'ประเภทผู้ประเมิน (user_type) — internal หรือ external',
    'ประเภทการประเมิน — ตนเอง (title มี "ประเมินตนเอง") หรือผู้อื่น',
    'ปีงบประมาณ (fiscal_year) — ปีเก่าเลือก eval เก่า, ปีปัจจุบันเลือก eval ใหม่',
]:
    doc.add_paragraph(f, style='List Bullet')

add_colored_heading('2.2 การจับกลุ่มระดับและองศา', level=2, color=RGBColor(0, 80, 120))
tbl2 = doc.add_table(rows=4, cols=4)
tbl2.style = 'Light Grid Accent 1'
for i, h in enumerate(['กลุ่ม', 'ระดับ', 'องศาที่รองรับ', 'หมายเหตุ']):
    set_cell(tbl2.rows[0].cells[i], h, bold=True, size=Pt(13))
data2 = [
    ('ผู้ว่าการ', '13', 'self, top, bottom, left, right', 'แบบเฉพาะปี 69+'),
    ('ผู้บริหาร', '9-12', 'self, top, bottom, left, right', 'ทุกองศา'),
    ('พนักงาน', '4-8', 'self, top, left', 'ไม่มี bottom/right'),
]
for ri, rd in enumerate(data2, 1):
    for ci, val in enumerate(rd):
        set_cell(tbl2.rows[ri].cells[ci], val, size=Pt(12))

add_colored_heading('2.3 ลำดับการเลือก', level=2, color=RGBColor(0, 80, 120))

doc.add_paragraph('กรณีประเมินผู้อื่น:', style='List Bullet')
for s in [
    '1) ใช้ evaluation_id จาก assignment (ถ้ามีและยัง published)',
    '2) Fallback: findByGrade(grade, user_type, fiscal_year)',
    '3) ปีเก่า → oldest(), ปีปัจจุบัน → latest()',
]:
    p = doc.add_paragraph(s)
    p.paragraph_format.left_indent = Cm(1.5)

doc.add_paragraph('กรณีประเมินตนเอง:', style='List Bullet')
for s in [
    '1) findSelfEvalByGrade(grade, fiscal_year) — หา title มี "ประเมินตนเอง"',
    '2) Fallback: ใช้แบบประเมินพนักงาน (internal) แทน',
]:
    p = doc.add_paragraph(s)
    p.paragraph_format.left_indent = Cm(1.5)

doc.add_paragraph()

# ══════════════════════════════════════════════
# Section 3: Detailed per-evaluation
# ══════════════════════════════════════════════
add_colored_heading('3. รายละเอียดคำถามแต่ละแบบประเมิน')

type_thai = {'rating': 'Rating 1-5', 'choice': 'ตัวเลือก', 'multiple_choice': 'เลือกหลายข้อ', 'open_text': 'ปลายเปิด'}

for idx, e in enumerate(evals):
    eid = e['id']
    total_q = sum(len(a['questions']) for p in e['parts'] for a in p['aspects'])
    fy = fy_map.get(eid, '?')

    doc.add_page_break()
    add_colored_heading(f"3.{idx+1} Eval {eid}: {e['title']}", level=2, color=RGBColor(0, 70, 130))

    # Info line
    info = doc.add_paragraph()
    r = info.add_run(f"ระดับ: {e['grade']}  |  ประเภท: {type_map.get(e['type'], e['type'])}  |  ปีงบ: 25{fy}  |  รวม: {total_q} คำถาม")
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(80, 80, 80)

    q_num = 0
    for part in e['parts']:
        part_q_count = sum(len(a['questions']) for a in part['aspects'])
        add_colored_heading(f"{part['title']} ({part_q_count} คำถาม)", level=3, color=RGBColor(0, 100, 60))

        for aspect in part['aspects']:
            if not aspect['questions']:
                continue

            # Aspect heading
            p_asp = doc.add_paragraph()
            r_asp = p_asp.add_run(f"■ {aspect['name']} ({len(aspect['questions'])} ข้อ)")
            r_asp.font.bold = True
            r_asp.font.size = Pt(13)
            r_asp.font.color.rgb = RGBColor(50, 50, 120)

            # Question table
            tbl_q = doc.add_table(rows=1, cols=4)
            tbl_q.style = 'Light List Accent 1'
            for i, h in enumerate(['ลำดับ', 'คำถาม', 'ประเภท', 'ตัวเลือก']):
                set_cell(tbl_q.rows[0].cells[i], h, bold=True, size=Pt(11), align=WD_ALIGN_PARAGRAPH.CENTER)

            # Set column widths
            for row_cells in [tbl_q.rows[0].cells]:
                row_cells[0].width = Cm(1.2)
                row_cells[1].width = Cm(10)
                row_cells[2].width = Cm(2)
                row_cells[3].width = Cm(4)

            for q in aspect['questions']:
                q_num += 1
                row = tbl_q.add_row().cells
                set_cell(row[0], str(q_num), size=Pt(11), align=WD_ALIGN_PARAGRAPH.CENTER)

                # Question title
                q_title = q['title']
                if len(q_title) > 120:
                    q_title = q_title[:117] + '...'
                set_cell(row[1], q_title, size=Pt(11))

                set_cell(row[2], type_thai.get(q['type'], q['type']), size=Pt(10), align=WD_ALIGN_PARAGRAPH.CENTER)

                # Options
                if q['options']:
                    opts_text = '\n'.join(q['options'][:6])
                    if len(q['options']) > 6:
                        opts_text += f'\n... +{len(q["options"])-6} more'
                    set_cell(row[3], opts_text, size=Pt(9))
                elif q['type'] == 'rating':
                    set_cell(row[3], '1-5 (น้อยที่สุด-มากที่สุด)', size=Pt(9))
                else:
                    set_cell(row[3], '-', size=Pt(9))

            doc.add_paragraph()  # spacing

# ══════════════════════════════════════════════
# Section 4: Comparison
# ══════════════════════════════════════════════
doc.add_page_break()
add_colored_heading('4. เปรียบเทียบแบบประเมินพนักงาน vs ประเมินตนเอง')

tbl_diff = doc.add_table(rows=6, cols=3)
tbl_diff.style = 'Light Grid Accent 1'
diff_data = [
    ['หัวข้อ', 'แบบประเมินพนักงาน (ภายใน/ภายนอก)', 'แบบประเมินตนเอง'],
    ['ส่วนที่ 1\n(Rating)', 'คำถาม Rating เหมือนกัน\nใช้สรรพนาม "ผู้บริหาร/พนักงาน"', 'คำถาม Rating เหมือนกัน\nใช้สรรพนาม "ท่าน" (สำหรับผู้ว่าการ)'],
    ['ส่วนที่ 2\n(วัฒนธรรมองค์กร)', '16 ข้อ Rating\n• 8 ข้อ ระดับยอมรับ/ความสำคัญ\n• 8 ข้อ ระดับแสดงพฤติกรรม', '8 ข้อ Choice/MC\n• 4 ข้อ การรับรู้ค่านิยม\n• 2 ข้อ ความเข้าใจค่านิยม\n• 2 ข้อ พฤติกรรมพึงประสงค์'],
    ['ส่วนที่ 3\n(ปลายเปิด)', 'ไม่มี (ยกเว้นผู้ว่าการ: 3 ข้อ)', 'มี 1 ข้อ (ยกตัวอย่างพฤติกรรม AQ)'],
    ['จำนวนรวม\n(9-12)', '39 ข้อ (ปี 68) / 46 ข้อ (ปี 69)', '32 ข้อ (ทั้ง 2 ปี)'],
    ['จำนวนรวม\n(4-8)', '29 ข้อ (ปี 68) / 39 ข้อ (ปี 69)', '22 ข้อ (ทั้ง 2 ปี)'],
]
for ri, rd in enumerate(diff_data):
    for ci, val in enumerate(rd):
        set_cell(tbl_diff.rows[ri].cells[ci], val, bold=(ri == 0), size=Pt(12 if ri == 0 else 11))

doc.add_paragraph()

# ══════════════════════════════════════════════
# Section 5: Changes 68→69
# ══════════════════════════════════════════════
add_colored_heading('5. การเปลี่ยนแปลงจากปี 2568 สู่ 2569')

changes = [
    'เพิ่มแบบประเมินผู้ว่าการ กนอ. แยกจากกลุ่ม 9-12 (Eval 33, 34, 38)',
    'ขยายช่วง Grade จาก 5-8 เป็น 4-8 — รองรับพนักงานระดับ 4',
    'ผู้บริหาร 9-12 ภายใน: เพิ่มด้าน "การบริหารความเสี่ยงและความรับผิดชอบ" + เพิ่มข้อในด้านอื่นๆ (23→30 ข้อ ส่วน 1)',
    'พนักงาน 4-8: เพิ่มด้าน "การคิดเชิงนวัตกรรม" และ "การเรียนรู้และปรับตัว" (13→23 ข้อ ส่วน 1)',
    'ผู้ว่าการภายใน: คำถามเฉพาะ "ผวก. กนอ." + 3 ข้อปลายเปิด ไม่มีส่วนวัฒนธรรมองค์กร Rating',
    'ประเมินตนเองทุกระดับ: ส่วนที่ 2 ยังคงเป็น 8 ข้อ Choice/MC เหมือนปี 68',
    'ระบบเลือกแบบประเมินตามปีงบประมาณ: ปีเก่า→eval เก่า, ปีปัจจุบัน→eval ใหม่',
]
for c in changes:
    doc.add_paragraph(c, style='List Bullet')

# ══════════════════════════════════════════════
# Save
# ══════════════════════════════════════════════
output_path = os.path.join(os.path.dirname(__file__), 'evaluation_forms_detail_analysis.docx')
doc.save(output_path)
print(f'Saved to: {output_path}')
