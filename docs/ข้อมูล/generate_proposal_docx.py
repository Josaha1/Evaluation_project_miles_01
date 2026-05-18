#!/usr/bin/env python3
"""Generate project_proposal.docx from project_proposal.md content."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, **kwargs):
    """Set cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", "4")}" w:space="0" '
            f'w:color="{val.get("color", "000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)

def add_table(doc, headers, rows, col_widths=None, header_color="1F4E79"):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = 'TH Sarabun New'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_color)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        is_bold_row = any(str(cell).startswith('**') or str(cell).startswith('รวม') for cell in row_data)
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            text = str(cell_text).replace('**', '')
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.font.size = Pt(10)
            run.font.name = 'TH Sarabun New'
            if is_bold_row:
                run.bold = True
                set_cell_shading(cell, "D6E4F0")
            elif r_idx % 2 == 1:
                set_cell_shading(cell, "F2F7FB")

    # Set column widths
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table

def add_heading_styled(doc, text, level=1):
    """Add a styled heading."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'TH Sarabun New'
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    return h

def add_paragraph_styled(doc, text, bold=False, italic=False, size=12, indent=False, color=None):
    """Add a styled paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'TH Sarabun New'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    return p

def add_bullet(doc, text, bold_prefix=None):
    """Add a bullet point."""
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.name = 'TH Sarabun New'
        run.font.size = Pt(12)
        run = p.add_run(text)
        run.font.name = 'TH Sarabun New'
        run.font.size = Pt(12)
    else:
        run = p.add_run(text)
        run.font.name = 'TH Sarabun New'
        run.font.size = Pt(12)
    return p

def add_code_block(doc, text):
    """Add a code block with gray background."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    # Add shading to paragraph
    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
    pPr.append(shading)
    return p

def add_info_box(doc, lines):
    """Add an info box (quoted block)."""
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        pPr = p._p.get_or_add_pPr()
        # Add left border
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:left w:val="single" w:sz="18" w:space="8" w:color="1F4E79"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E8F0FE" w:val="clear"/>')
        pPr.append(shading)
        # Parse bold parts
        parts = line.split('**')
        for i, part in enumerate(parts):
            if not part:
                continue
            run = p.add_run(part)
            run.font.name = 'TH Sarabun New'
            run.font.size = Pt(11)
            if i % 2 == 1:  # odd indices are bold
                run.bold = True

def generate_docx():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'TH Sarabun New'
    font.size = Pt(12)

    # ===== COVER / TITLE =====
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ข้อเสนอโครงการพัฒนาระบบการประเมิน 360 องศา กนอ.")
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = 'TH Sarabun New'
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Project Proposal — 360-Degree Evaluation System (Phase 3)")
    run.font.size = Pt(16)
    run.font.name = 'TH Sarabun New'
    run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)

    doc.add_paragraph()

    # Info box
    add_info_box(doc, [
        "**วันที่จัดทำ**: 1 มีนาคม 2569",
        "**ระยะงวด**: งวดงานที่ 3 (180 วัน)",
        "**กำหนดส่งมอบ**: 14 มิถุนายน 2569",
        "**ระบบ**: Laravel + React (Inertia.js) + MariaDB บน Hostinger",
    ])

    doc.add_page_break()

    # ===== 1. ภาพรวมโครงการ =====
    add_heading_styled(doc, "1. ภาพรวมโครงการ", level=1)
    add_paragraph_styled(doc,
        "ระบบการประเมิน 360 องศา กนอ. (การนิคมอุตสาหกรรมแห่งประเทศไทย) เป็นระบบประเมินบุคลากร"
        "ครอบคลุมทุกระดับ ตั้งแต่ผู้ว่าการ ผู้บริหาร ไปจนถึงพนักงานระดับ 4-12 โดยประเมินจากหลาย"
        "มุมมอง (360 องศา) ได้แก่ ผู้บังคับบัญชา ผู้ใต้บังคับบัญชา เพื่อนร่วมงาน องค์กรภายนอก และการประเมินตนเอง"
    )

    add_heading_styled(doc, "ระบบที่มีอยู่แล้ว (Baseline)", level=2)
    add_table(doc,
        ["รายการ", "รายละเอียด"],
        [
            ["Framework", "Laravel + React (Inertia.js)"],
            ["Database", "MariaDB 10.4 — milesconsultdb"],
            ["Hosting", "Hostinger (12GB RAM, 6 Core, 300 PHP Workers)"],
            ["ข้อมูลที่มีอยู่", "~207,000+ คำตอบ, 1,200+ users, 8,200+ assignments"],
            ["แบบประเมินปัจจุบัน", "5 ชุด (ผู้บริหาร 9-12, พนักงาน 4-8, ประเมินตนเอง)"],
        ],
        col_widths=[5, 11]
    )

    # ===== 2. Timeline =====
    doc.add_paragraph()
    add_heading_styled(doc, "2. Timeline งวดงานที่ 3", level=1)

    add_code_block(doc,
        "1 มี.ค. 2569  ──────────────────────────────────  14 มิ.ย. 2569\n"
        "    │                                                      │\n"
        "    ├── Sprint 0 ──┤ Sprint 1 ────────┤ Sprint 2 ──┤      │\n"
        "    │  1–6 มี.ค.   │  7–20 มี.ค.      │  21–31มี.ค │      │\n"
        "    │  (6 วัน)     │  (14 วัน)        │  (11 วัน)  │      │\n"
        "    ▼              ▼                  ▼             ▼      ▼\n"
        "  Deploy        Task 3+4           Task 5        Live   Report\n"
        "  Task 1,2     External+UX        Dashboard   (1เม.ย.)   Due"
    )

    add_heading_styled(doc, "Milestone หลัก", level=2)
    add_table(doc,
        ["วันที่", "Milestone", "สถานะที่ต้องการ"],
        [
            ["2–6 มี.ค. 2569", "ทดสอบความเที่ยงตรง + ชี้แจงระบบ", "ระบบ stable, ผู้ว่าการ eval พร้อม"],
            ["7–31 มี.ค. 2569", "พัฒนาและปรับปรุงระบบ", "External login, UX, Dashboard ครบ"],
            ["1 เม.ย. 2569", "เริ่ม Live Evaluation", "Production-ready 100%"],
            ["31 พ.ค. 2569", "สิ้นสุด Live Evaluation", "ข้อมูลครบ, ไม่มี critical bugs"],
            ["1–14 มิ.ย. 2569", "Progress Report 2", "Export/Report สมบูรณ์"],
        ],
        col_widths=[4, 5.5, 6.5]
    )

    # ===== 3. ขอบเขตงาน =====
    doc.add_page_break()
    add_heading_styled(doc, "3. ขอบเขตงาน (Scope of Work)", level=1)

    # --- Task 1 ---
    add_heading_styled(doc, "Task 1 — แบบประเมินผู้ว่าการ กนอ.", level=2)
    add_paragraph_styled(doc, "วัตถุประสงค์: สร้างแบบประเมิน 360 องศาสำหรับผู้ว่าการ กนอ. ซึ่งยังไม่มีในระบบปัจจุบัน", bold=True)

    add_heading_styled(doc, "รายละเอียดงาน", level=3)
    add_table(doc,
        ["รายการ", "รายละเอียด"],
        [
            ["แบบประเมิน Internal", "1 ชุด — ผู้ประเมินภายใน (ผู้ใต้บังคับบัญชา, เพื่อนร่วมงาน)"],
            ["แบบประเมิน External", "1 ชุด — ผู้ประเมินภายนอก (องค์กรภายนอก, คู่ค้า)"],
            ["หมวดหมู่ (Aspects)", "6 หมวด ตามเกณฑ์ กนอ."],
            ["โครงสร้าง", "evaluation → parts → aspects → questions → options"],
        ],
        col_widths=[5, 11]
    )

    doc.add_paragraph()
    add_heading_styled(doc, "งานที่ต้องทำ", level=3)
    add_bullet(doc, "สร้าง Database Migration สำหรับ evaluations, parts, aspects, questions ผู้ว่าการ")
    add_bullet(doc, "Seed ข้อมูลคำถาม + ตัวเลือกทั้ง 2 ชุด (internal + external)")
    add_bullet(doc, "Admin UI: จัดการแบบประเมินผู้ว่าการผ่าน Admin panel")
    add_bullet(doc, "Assignment System: กำหนดว่าใครประเมินผู้ว่าการได้บ้าง")
    add_bullet(doc, "ทดสอบ end-to-end flow")

    add_heading_styled(doc, "Angle Weights สำหรับผู้ว่าการ", level=3)
    add_table(doc,
        ["Angle", "น้ำหนัก"],
        [
            ["ผู้บังคับบัญชา (top)", "25%"],
            ["ผู้ใต้บังคับบัญชา (bottom)", "25%"],
            ["ประเมินตนเอง (self)", "10%"],
            ["เพื่อนร่วมงาน (left)", "20%"],
            ["องค์กรภายนอก (right)", "20%"],
        ],
        col_widths=[8, 8]
    )
    add_paragraph_styled(doc, "ประมาณชั่วโมงทำงาน: 16–24 ชั่วโมง", bold=True)

    # --- Task 2 ---
    doc.add_paragraph()
    add_heading_styled(doc, "Task 2 — Maintenance Agreement (MA) ระบบรายเดือน", level=2)
    add_paragraph_styled(doc, "วัตถุประสงค์: ดูแลระบบให้ stable ตลอดช่วง Live Evaluation (เม.ย.–พ.ค.)", bold=True)

    add_heading_styled(doc, "บริการที่ครอบคลุม", level=3)
    add_table(doc,
        ["บริการ", "ความถี่"],
        [
            ["Monitor Hostinger (CPU, RAM, PHP Workers)", "ทุกวัน"],
            ["Database Backup", "ทุกสัปดาห์"],
            ["Security Patches (Laravel, dependencies)", "เมื่อมี update สำคัญ"],
            ["Performance Tuning (query optimization, cache)", "เมื่อจำเป็น"],
            ["Bug Fix (non-critical)", "ภายใน 48 ชั่วโมง"],
            ["Bug Fix (critical — ระบบใช้ไม่ได้)", "ภายใน 4 ชั่วโมง"],
            ["Hostinger Config Optimization", "ครั้งแรก one-time setup"],
        ],
        col_widths=[9, 7]
    )

    add_heading_styled(doc, "Hostinger Server Specs (ที่มีอยู่)", level=3)
    add_code_block(doc,
        "Disk Space  : 300 GB        เพียงพอมาก\n"
        "RAM         : 12,288 MB     รองรับ Laravel + DB + Queue ได้ดี\n"
        "CPU Cores   : 6             รองรับ concurrent users สูง\n"
        "PHP Workers : 300           ไม่ติด bottleneck\n"
        "Max Process : 600           Queue jobs รองรับได้\n"
        "Bandwidth   : Unlimited     ไม่ต้องห่วงเรื่อง traffic"
    )
    add_paragraph_styled(doc, "ประมาณชั่วโมงทำงาน: 8–16 ชั่วโมง/เดือน", bold=True)

    # --- Task 3 ---
    doc.add_page_break()
    add_heading_styled(doc, "Task 3 — ระบบ External Organization Login + องศาขวา", level=2)
    add_paragraph_styled(doc,
        "วัตถุประสงค์: สร้างระบบให้องค์กรภายนอก (ไม่มีรหัสพนักงาน) สามารถเข้ามาประเมิน"
        "ผู้บริหาร/ผู้ว่าการ กนอ. ได้ พร้อมระบุได้ว่าผลมาจากองค์กรใด", bold=True)

    add_heading_styled(doc, "การตัดสินใจ: Access Code + QR Code Hybrid", level=3)
    add_table(doc,
        ["เหตุผล", "รายละเอียด"],
        [
            ["ไม่พึ่ง 3rd Party", "ไม่ต้องใช้ LINE API, SMS Gateway, SMTP"],
            ["แจกได้ทุกช่องทาง", "พิมพ์กระดาษ / LINE / SMS / งานประชุม"],
            ["Tag Organization ชัดเจน", "รหัสองค์กรฝังใน Code IEAT-[ORG]-XXXXXX"],
            ["Report แยกตามองค์กร", "รู้ว่าคะแนนมาจาก บริษัท A หรือ B"],
            ["Revoke ได้ทันที", "ถ้า code หลุด Admin ยกเลิกได้เลย"],
            ["Dev เร็ว", "ไม่มี OAuth flow ซับซ้อน ทันกำหนด Sprint 0"],
        ],
        col_widths=[5, 11]
    )

    add_heading_styled(doc, "รูปแบบ Access Code", level=3)
    add_code_block(doc,
        "Access Code : IEAT-[ORG_CODE]-[RANDOM_6]\n"
        "ตัวอย่าง   : IEAT-BKKP-A7X3K2\n\n"
        "  IEAT    = ชื่อระบบ (คงที่)\n"
        "  BKKP    = รหัสย่อองค์กร (Admin กำหนด)\n"
        "  A7X3K2  = รหัสเฉพาะบุคคล 6 ตัว (auto-generate)\n\n"
        "QR Code    : https://[domain]/external/evaluate?token=IEAT-BKKP-A7X3K2"
    )

    add_heading_styled(doc, "รายละเอียดงานที่ต้องทำ", level=3)
    add_table(doc,
        ["รายการ", "รายละเอียด", "ชั่วโมง"],
        [
            ["Database Migration", "3 ตารางใหม่ + ALTER answers", "4"],
            ["ExternalAuth Middleware", "Session validation + route protection", "4"],
            ["ExternalEvaluatorController", "Login, QR entry, Confirm, Dashboard, Evaluate, Submit", "16"],
            ["Admin: External Organizations", "CRUD + org_code management", "8"],
            ["Admin: Access Codes", "CRUD + Generate + Revoke + Regenerate", "12"],
            ["QR Code Generation", "simplesoftwareio/simple-qrcode library", "4"],
            ["Print Card PDF", "พิมพ์บัตรแจก", "4"],
            ["React Pages (External)", "Login, Confirm, Dashboard, Evaluate", "14"],
            ["React Pages (Admin)", "Org management, Code management", "10"],
            ["Report Integration", "คะแนนแยกตามองค์กรใน AdminDashboard", "8"],
            ["Security", "Rate limiting, session expiry, audit log", "4"],
            ["Testing", "End-to-end ทุก flow", "8"],
            ["**รวม**", "", "**96 ชั่วโมง**"],
        ],
        col_widths=[5, 7.5, 3.5]
    )
    add_paragraph_styled(doc, "ประมาณชั่วโมงทำงาน: 80–96 ชั่วโมง", bold=True)

    # --- Task 4 ---
    doc.add_page_break()
    add_heading_styled(doc, "Task 4 — Adjust Workflow ระบบเก่า", level=2)
    add_paragraph_styled(doc, "วัตถุประสงค์: ปรับ workflow การใช้งานให้เข้าใจง่ายขึ้น โดยไม่กระทบข้อมูลเก่า", bold=True)

    add_heading_styled(doc, "ปัญหาปัจจุบัน vs. การแก้ไข", level=3)
    add_table(doc,
        ["#", "ปัญหาปัจจุบัน", "การแก้ไข"],
        [
            ["1", "Dashboard ไม่ชัดว่าต้องทำอะไรก่อน-หลัง", "Step-based UI (Step 1: ตนเอง → Step 2: คนอื่น)"],
            ["2", "ไม่รู้ว่า auto-save ทำงานหรือยัง", "Save indicator + confirmation toast"],
            ["3", "ส่งแบบประเมินเสร็จแล้วไม่มี feedback ชัดเจน", "Modal confirmation + summary page"],
            ["4", "Admin ต้อง assign ทีละคน (ช้า)", "Bulk assignment: เลือกหลายคนพร้อมกัน"],
            ["5", "ไม่เห็นภาพรวม progress ของตัวเอง", "Progress tracker ครบทุก assignment"],
            ["6", "หน้าจอซับซ้อน มีข้อมูลเยอะเกินไป", "Clean layout, ซ่อนข้อมูลรอง"],
        ],
        col_widths=[1.5, 6, 8.5]
    )
    add_paragraph_styled(doc, "ประมาณชั่วโมงทำงาน: 32–48 ชั่วโมง", bold=True)

    # --- Task 5 ---
    doc.add_paragraph()
    add_heading_styled(doc, "Task 5 — AdminDashboard ครบทุกระบบ", level=2)
    add_paragraph_styled(doc, "วัตถุประสงค์: ให้ AdminDashboard รองรับทุก evaluation type แบบ dynamic ไม่ hardcode eval IDs", bold=True)

    add_heading_styled(doc, "งานที่ต้องทำ", level=3)
    add_table(doc,
        ["รายการ", "ชั่วโมง"],
        [
            ["Refactor Controller: ลบ hardcoded IDs ให้ dynamic", "12"],
            ["KPI Cards รองรับทุก eval type", "10"],
            ["Charts: แยก view ตาม grade group", "14"],
            ["Weighted Score: ถูกต้องทุก grade level", "10"],
            ["External Org section ใน Report", "12"],
            ["Export: ครบทุก eval type รวม ผู้ว่าการ + external", "14"],
            ["Individual Report: ทุก grade", "8"],
            ["UI Polish + Filter UX", "10"],
            ["Testing", "8"],
            ["**รวม**", "**98 ชั่วโมง**"],
        ],
        col_widths=[12, 4]
    )
    add_paragraph_styled(doc, "ประมาณชั่วโมงทำงาน: 80–100 ชั่วโมง", bold=True)

    # ===== 4. แผนการทำงาน =====
    doc.add_page_break()
    add_heading_styled(doc, "4. แผนการทำงาน — Developer 1 คน (OT Model)", level=1)

    add_info_box(doc, [
        "**รูปแบบ**: Developer 1 คน ทำงานนอกเวลา (OT) ควบคู่กับงานประจำ",
        "**เงินเดือนปกติ**: 40,000 บาท/เดือน",
        "**ทำงานวันธรรมดาหลังเลิกงาน**: ~2–3 ชม./วัน",
        "**ทำงานวันหยุด (เสาร์-อาทิตย์)**: ~6–8 ชม./วัน",
        "**กำลังการผลิต OT ต่อสัปดาห์**: ~25–30 ชม.",
    ])

    add_heading_styled(doc, "เหตุผลที่ลดชั่วโมงจากแผนเดิม (~28%)", level=2)
    add_table(doc,
        ["#", "เหตุผล", "ผลกระทบ"],
        [
            ["1", "ใช้ AI-assisted development (Claude Code) เร่งงาน coding, testing, docs", "ลด 20–30% ต่อ task"],
            ["2", "Developer คนเดียว — ไม่มี communication overhead, merge conflict", "ลด 5–10%"],
            ["3", "Reuse pattern จากระบบที่มีอยู่ (เช่น CRUD org structure ที่สร้างแล้ว)", "ลด 10–15%"],
            ["4", "ลำดับงานต่อเนื่อง — ไม่ต้อง context switch ระหว่าง 2 คน", "ลด 5%"],
        ],
        col_widths=[1.5, 10, 4.5]
    )

    # Sprint 0
    doc.add_paragraph()
    add_heading_styled(doc, "Sprint 0 — เร่งด่วน (1 มี.ค. – 6 มี.ค.) = 6 วัน | 22 ชม.", level=2)
    add_paragraph_styled(doc, "เป้าหมาย: ระบบพร้อมทดสอบ + ผู้ว่าการ eval ใช้งานได้", bold=True)
    add_table(doc,
        ["สัปดาห์", "งาน", "ชม. OT"],
        [
            ["สัปดาห์ 1 (จ–ศ)", "สร้าง Migration ผู้ว่าการ + Seed คำถาม + Admin UI", "12"],
            ["สัปดาห์ 1 (ส–อา)", "Deploy + Config Hostinger + ทดสอบ Assignment", "10"],
            ["**รวม**", "", "**22**"],
        ],
        col_widths=[4, 9, 3]
    )

    # Sprint 1
    doc.add_paragraph()
    add_heading_styled(doc, "Sprint 1 — พัฒนาหลัก (7 มี.ค. – 6 เม.ย.) = 31 วัน | 92 ชม.", level=2)
    add_paragraph_styled(doc, "เป้าหมาย: External Access Code + Workflow ใหม่", bold=True)
    add_table(doc,
        ["สัปดาห์", "งาน", "ชม. OT"],
        [
            ["สัปดาห์ 1", "External Organizations DB + Migration + Middleware", "22"],
            ["สัปดาห์ 2", "Admin: จัดการ Org + Access Codes + QR Generate", "22"],
            ["สัปดาห์ 3", "External Login/Confirm/Dashboard + Evaluation Form", "22"],
            ["สัปดาห์ 4", "User Dashboard Workflow redesign + Bulk Assignment", "20"],
            ["สัปดาห์ 4.5", "Testing + Bug fixes", "6"],
            ["**รวม**", "", "**92**"],
        ],
        col_widths=[4, 9, 3]
    )

    # Sprint 2
    doc.add_paragraph()
    add_heading_styled(doc, "Sprint 2 — AdminDashboard ครบระบบ (7 เม.ย. – 30 เม.ย.) = 24 วัน | 64 ชม.", level=2)
    add_paragraph_styled(doc, "เป้าหมาย: Dashboard ครบ + Production-ready", bold=True)
    add_table(doc,
        ["สัปดาห์", "งาน", "ชม. OT"],
        [
            ["สัปดาห์ 1", "Refactor Controller (dynamic) + KPI Cards ทุก eval", "20"],
            ["สัปดาห์ 2", "Weighted Score ครบ + Charts + Filter", "20"],
            ["สัปดาห์ 3", "Export ครบทุก eval type + Individual Report", "18"],
            ["สัปดาห์ 3.5", "Security + Integration testing + Deploy", "6"],
            ["**รวม**", "", "**64**"],
        ],
        col_widths=[4, 9, 3]
    )

    # Sprint 3
    doc.add_paragraph()
    add_heading_styled(doc, "Sprint 3 — Live Evaluation + MA (1 พ.ค. – 31 พ.ค.) = 31 วัน | 10 ชม./เดือน", level=2)
    add_paragraph_styled(doc, "เป้าหมาย: ระบบ stable ตลอดช่วง Live", bold=True)
    add_table(doc,
        ["รายการ", "ความถี่"],
        [
            ["Monitor Hostinger (CPU, RAM, PHP Workers)", "ทุกวัน (5 นาที)"],
            ["Database Backup verify", "ทุกสัปดาห์"],
            ["Critical bug fix", "ภายใน 4 ชั่วโมง"],
            ["Non-critical bug fix", "ภายใน 48 ชั่วโมง"],
        ],
        col_widths=[10, 6]
    )

    # Sprint 4
    doc.add_paragraph()
    add_heading_styled(doc, "Sprint 4 — Report Delivery (1–14 มิ.ย.) = 14 วัน | 14 ชม.", level=2)
    add_table(doc,
        ["รายการ", "ชม."],
        [
            ["Export ทดสอบกับ Live data จริงครบ", "4"],
            ["Dashboard Summary สำหรับ Progress Report 2", "4"],
            ["PDF/Excel รายงานทุกกลุ่ม", "4"],
            ["UAT + ส่งมอบ", "2"],
            ["**รวม**", "**14**"],
        ],
        col_widths=[12, 4]
    )

    # ===== 5. สรุปชั่วโมง =====
    doc.add_page_break()
    add_heading_styled(doc, "5. สรุปชั่วโมงทำงาน (ปรับลดสำหรับ 1 คน + AI-assisted)", level=1)

    add_heading_styled(doc, "เปรียบเทียบชั่วโมงเดิม vs. ใหม่", level=2)
    add_table(doc,
        ["Task", "งาน", "เดิม (2 คน)", "ปรับลด (1 คน)", "ลด%"],
        [
            ["Task 1", "แบบประเมินผู้ว่าการ", "20 ชม.", "14 ชม.", "-30%"],
            ["Task 3", "External Access Code + QR", "88 ชม.", "64 ชม.", "-27%"],
            ["Task 4", "Workflow Adjustment", "40 ชม.", "28 ชม.", "-30%"],
            ["Task 5", "AdminDashboard ครบระบบ", "90 ชม.", "64 ชม.", "-29%"],
            ["Deploy", "Hostinger Setup + Config", "10 ชม.", "8 ชม.", "-20%"],
            ["Sprint 4", "Report Preparation", "20 ชม.", "14 ชม.", "-30%"],
            ["**รวม One-time**", "", "**268 ชม.**", "**192 ชม.**", "**-28%**"],
            ["Task 2", "MA รายเดือน", "12 ชม.", "10 ชม.", "-17%"],
        ],
        col_widths=[2.5, 4.5, 3, 3, 3],
        header_color="2E75B6"
    )

    doc.add_paragraph()
    add_heading_styled(doc, "กำลังการผลิต OT vs. ความต้องการ", level=2)
    add_table(doc,
        ["Sprint", "ระยะเวลา", "ชม. ที่ต้องการ", "ชม. OT ที่มี", "สถานะ"],
        [
            ["Sprint 0", "6 วัน (~1 สัปดาห์)", "22", "~55", "เพียงพอ"],
            ["Sprint 1", "31 วัน (~4.5 สัปดาห์)", "92", "~125", "เพียงพอ"],
            ["Sprint 2", "24 วัน (~3.5 สัปดาห์)", "64", "~95", "เพียงพอ"],
            ["Sprint 3", "31 วัน (MA)", "10", "~30", "เพียงพอ"],
            ["Sprint 4", "14 วัน (~2 สัปดาห์)", "14", "~55", "เพียงพอ"],
            ["**รวม**", "**106 วัน**", "**202**", "**~360**", "**มี buffer 78%**"],
        ],
        col_widths=[2.5, 4.5, 3, 3, 3]
    )

    # ===== 6. ประเมินราคา =====
    doc.add_page_break()
    add_heading_styled(doc, "6. ประเมินราคา — OT Model (Developer 1 คน)", level=1)

    add_heading_styled(doc, "ฐานคำนวณ OT ตาม พ.ร.บ.คุ้มครองแรงงาน", level=2)
    add_table(doc,
        ["รายการ", "การคำนวณ", "อัตรา"],
        [
            ["เงินเดือน", "—", "40,000 บาท/เดือน"],
            ["ค่าจ้างรายวัน", "40,000 / 30", "1,333.33 บาท/วัน"],
            ["ค่าจ้างรายชั่วโมง", "1,333.33 / 8", "166.67 บาท/ชม."],
            ["OT วันธรรมดา", "166.67 x 1.5", "250 บาท/ชม."],
            ["ทำงานวันหยุด (ชม.ปกติ)", "166.67 x 1.0 (เพิ่มเติม)", "167 บาท/ชม."],
            ["OT วันหยุด (เกิน 8 ชม.)", "166.67 x 3.0", "500 บาท/ชม."],
        ],
        col_widths=[5.5, 5.5, 5]
    )

    add_info_box(doc, [
        "**หมายเหตุ**: ลูกจ้างรายเดือนทำงานวันหยุด ได้ค่าจ้างเพิ่ม 1 เท่า (มาตรา 62)",
        "ล่วงเวลาวันหยุด ได้ 3 เท่า (มาตรา 63)",
    ])

    doc.add_paragraph()
    add_heading_styled(doc, "ประมาณสัดส่วน OT แยกประเภท", level=2)
    add_table(doc,
        ["Sprint", "รวม (ชม.)", "OT วันธรรมดา (1.5x)", "วันหยุด ปกติ (1x)", "OT วันหยุด (3x)"],
        [
            ["Sprint 0", "22", "12", "8", "2"],
            ["Sprint 1", "92", "50", "28", "14"],
            ["Sprint 2", "64", "40", "18", "6"],
            ["Sprint 4", "14", "10", "4", "0"],
            ["**รวม Dev**", "**192**", "**112**", "**58**", "**22**"],
            ["MA (4 เดือน)", "40", "32", "8", "0"],
            ["**รวมทั้งหมด**", "**232**", "**144**", "**66**", "**22**"],
        ],
        col_widths=[3, 3, 3, 3.5, 3.5],
        header_color="2E75B6"
    )

    doc.add_paragraph()
    add_heading_styled(doc, "คำนวณค่าตอบแทน OT", level=2)

    add_heading_styled(doc, "งานพัฒนา (One-time) — 192 ชั่วโมง", level=3)
    add_table(doc,
        ["ประเภท OT", "ชั่วโมง", "อัตรา (บาท/ชม.)", "ค่าตอบแทน (บาท)"],
        [
            ["OT วันธรรมดา (x1.5)", "112", "250", "28,000"],
            ["ทำงานวันหยุด ปกติ (x1.0 เพิ่ม)", "58", "167", "9,686"],
            ["OT วันหยุด (x3.0)", "22", "500", "11,000"],
            ["**รวม One-time**", "**192**", "", "**48,686**"],
        ],
        col_widths=[5.5, 3, 3.5, 4],
        header_color="C55A11"
    )

    doc.add_paragraph()
    add_heading_styled(doc, "MA — 40 ชั่วโมง (4 เดือน x 10 ชม.)", level=3)
    add_table(doc,
        ["ประเภท OT", "ชั่วโมง", "อัตรา (บาท/ชม.)", "ค่าตอบแทน (บาท)"],
        [
            ["OT วันธรรมดา (x1.5)", "32", "250", "8,000"],
            ["ทำงานวันหยุด ปกติ (x1.0 เพิ่ม)", "8", "167", "1,336"],
            ["**รวม MA**", "**40**", "", "**9,336**"],
        ],
        col_widths=[5.5, 3, 3.5, 4],
        header_color="C55A11"
    )

    # Summary cost
    doc.add_paragraph()
    add_heading_styled(doc, "สรุปราคารวม Project (งวดงานที่ 3) — OT Model", level=2)
    add_table(doc,
        ["รายการ", "ชั่วโมง", "ค่าตอบแทน (บาท)"],
        [
            ["One-time Development", "192 ชม.", "48,686"],
            ["MA 4 เดือน (มี.ค.–มิ.ย.)", "40 ชม.", "9,336"],
            ["**รวมทั้งสิ้น**", "**232 ชม.**", "**58,022**"],
            ["**ปัดเป็นเลขกลม**", "", "**~58,000 บาท**"],
        ],
        col_widths=[6, 5, 5],
        header_color="548235"
    )

    # Comparison
    doc.add_paragraph()
    add_heading_styled(doc, "เปรียบเทียบกับแผนเดิม (จ้างภายนอก 2 คน)", level=2)
    add_table(doc,
        ["รายการ", "แผนเดิม (จ้างนอก 2 คน)", "แผนใหม่ (OT 1 คน)", "ประหยัด"],
        [
            ["ชั่วโมงทำงาน", "268 + MA 48 = 316 ชม.", "192 + MA 40 = 232 ชม.", "-27%"],
            ["ค่าพัฒนา One-time", "214,400 บาท", "48,686 บาท", "-77%"],
            ["ค่า MA (4 เดือน)", "~32,000 บาท", "9,336 บาท", "-71%"],
            ["**รวมทั้งสิ้น**", "**~246,400 บาท**", "**~58,000 บาท**", "**-76%**"],
        ],
        col_widths=[4, 4.5, 4.5, 3]
    )

    add_info_box(doc, [
        "**ข้อแลกเปลี่ยน**: ใช้เวลานานขึ้น (Sprint 1 ขยายจาก 14 วัน → 31 วัน)",
        "แต่ยังอยู่ในกรอบเวลารวม 106 วัน (14 มิ.ย. 2569)",
    ])

    # Monthly payment
    doc.add_paragraph()
    add_heading_styled(doc, "การจ่ายค่า OT รายเดือน (ประมาณ)", level=2)
    add_table(doc,
        ["เดือน", "Sprint", "ชม. OT", "ค่า OT โดยประมาณ"],
        [
            ["มี.ค. 2569", "Sprint 0 + Sprint 1 (บางส่วน)", "70", "~18,200"],
            ["เม.ย. 2569", "Sprint 1 (ส่วนที่เหลือ) + Sprint 2 (เริ่ม)", "80", "~20,600"],
            ["พ.ค. 2569", "Sprint 2 (จบ) + Sprint 3 (MA)", "42", "~11,200"],
            ["มิ.ย. 2569", "Sprint 3 (MA) + Sprint 4", "40", "~8,000"],
            ["**รวม**", "", "**232**", "**~58,000**"],
        ],
        col_widths=[3.5, 6, 3, 3.5]
    )

    # ===== 7. Deliverables =====
    doc.add_page_break()
    add_heading_styled(doc, "7. Deliverables (สิ่งที่ส่งมอบ)", level=1)
    add_table(doc,
        ["#", "Deliverable", "กำหนดส่ง"],
        [
            ["D1", "แบบประเมินผู้ว่าการ กนอ. (internal + external) บน production", "1 มี.ค."],
            ["D2", "ระบบ External Organization Login + ประเมินองศาขวา", "20 มี.ค."],
            ["D3", "Workflow ใหม่ (User Dashboard + Admin Bulk Assign)", "20 มี.ค."],
            ["D4", "AdminDashboard รองรับทุก eval type + Export ครบ", "31 มี.ค."],
            ["D5", "Production Server Config + Load Test Report", "31 มี.ค."],
            ["D6", "MA Report (เม.ย.–พ.ค.) — สรุปการดูแลระบบรายเดือน", "ทุกสิ้นเดือน"],
            ["D7", "Export/Report สำหรับ Progress Report 2", "10 มิ.ย."],
        ],
        col_widths=[1.5, 10.5, 4]
    )

    # ===== 8. ข้อกำหนด =====
    doc.add_paragraph()
    add_heading_styled(doc, "8. ข้อกำหนดและเงื่อนไข", level=1)

    add_heading_styled(doc, "สิ่งที่ผู้ว่าจ้างต้องจัดเตรียม", level=2)
    add_bullet(doc, "ข้อมูลคำถามแบบประเมินผู้ว่าการ กนอ. (ฉบับสมบูรณ์)")
    add_bullet(doc, "รายชื่อองค์กรภายนอกที่ต้องการเชิญประเมิน")
    add_bullet(doc, "Hostinger credentials สำหรับ deploy")
    add_bullet(doc, "Email server / SMTP credentials สำหรับส่ง invitation")

    add_heading_styled(doc, "ข้อยกเว้นจาก Scope", level=2)
    add_bullet(doc, "การเพิ่มคำถามใหม่ในแบบประเมินที่มีอยู่แล้ว (ทำเองผ่าน Admin)")
    add_bullet(doc, "การออกแบบ Graphics / Branding ใหม่")
    add_bullet(doc, "การแปลภาษาอื่นนอกจากภาษาไทย")
    add_bullet(doc, "การ integrate กับระบบ HR ภายนอก (ถ้ามี ต้องประเมินเพิ่ม)")

    add_heading_styled(doc, "SLA (Service Level Agreement) — ช่วง Live Evaluation", level=2)
    add_table(doc,
        ["ประเภท", "ตอบกลับ", "แก้ไข"],
        [
            ["Critical (ระบบล่ม, login ไม่ได้, data loss)", "2 ชั่วโมง", "4 ชั่วโมง"],
            ["High (ฟีเจอร์หลักใช้ไม่ได้)", "4 ชั่วโมง", "24 ชั่วโมง"],
            ["Medium (ฟีเจอร์รองมีปัญหา)", "8 ชั่วโมง", "48 ชั่วโมง"],
            ["Low (UI เล็กน้อย, ปรับปรุง)", "24 ชั่วโมง", "Sprint ถัดไป"],
        ],
        col_widths=[7, 4.5, 4.5]
    )

    # ===== 9. สรุป =====
    doc.add_paragraph()
    add_heading_styled(doc, "9. สรุป", level=1)
    add_table(doc,
        ["หัวข้อ", "รายละเอียด"],
        [
            ["โครงการ", "พัฒนาระบบการประเมิน 360 องศา กนอ. (Phase 3)"],
            ["ระยะเวลา", "1 มี.ค. – 14 มิ.ย. 2569 (106 วัน)"],
            ["ทีมพัฒนา", "Developer 1 คน (ทำงานนอกเวลา OT) + AI-assisted"],
            ["เงินเดือนฐาน", "40,000 บาท/เดือน"],
            ["Tasks หลัก", "5 tasks + MA"],
            ["ชั่วโมง OT รวม", "232 ชั่วโมง (192 dev + 40 MA)"],
            ["ค่า OT รวม Project", "~58,000 บาท"],
            ["ค่า OT เฉลี่ยต่อเดือน", "~14,500 บาท/เดือน (4 เดือน)"],
            ["Authentication Approach", "Account-based + External Organizations Table"],
            ["Critical Deadline", "6 มีนาคม 2569 (ทดสอบระบบ)"],
        ],
        col_widths=[5, 11],
        header_color="548235"
    )

    # ===== 10. งานที่ดำเนินการแล้ว =====
    doc.add_page_break()
    add_heading_styled(doc, "10. งานที่ดำเนินการแล้ว (Completed Work)", level=1)

    add_heading_styled(doc, "10.1 ระบบจัดการโครงสร้างองค์กร (Organizational Structure Management)", level=2)
    add_paragraph_styled(doc, "สถานะ: Completed (23 ก.พ. 2569)", bold=True, color=RGBColor(0x00, 0x80, 0x00))

    add_paragraph_styled(doc, "ระบบ CRUD สำหรับจัดการโครงสร้างองค์กรทั้ง 4 entities ผ่าน Admin Panel:")
    add_table(doc,
        ["Entity", "Controller", "หน้า React", "Routes"],
        [
            ["สายงาน (Division)", "AdminDivisionController", "AdminDivisionIndex + Form", "6 routes"],
            ["หน่วยงาน (Department)", "AdminDepartmentController", "AdminDepartmentIndex + Form", "6 routes"],
            ["ตำแหน่ง (Position)", "AdminPositionController", "AdminPositionIndex + Form", "6 routes"],
            ["ฝ่าย (Faction)", "AdminFactionController", "AdminFactionIndex + Form", "6 routes"],
        ],
        col_widths=[4, 5, 4.5, 2.5]
    )

    doc.add_paragraph()
    add_paragraph_styled(doc, "ฟีเจอร์ที่ครอบคลุม:", bold=True)
    add_bullet(doc, "CRUD ครบ (สร้าง, ดู, แก้ไข, ลบ) ทุก entity")
    add_bullet(doc, "ค้นหา + Pagination (10 รายการ/หน้า)")
    add_bullet(doc, "Filter by parent entity")
    add_bullet(doc, "Data integrity protection (ป้องกันลบ entity ที่มีสมาชิก)")
    add_bullet(doc, "Dark mode support + Breadcrumb navigation + Toast notifications")

    add_heading_styled(doc, "10.2 ระบบรายงานและ Export (Report & Export System)", level=2)
    add_paragraph_styled(doc, "สถานะ: Completed", bold=True, color=RGBColor(0x00, 0x80, 0x00))
    add_bullet(doc, "AdminEvaluationReport.tsx — Dashboard, Analytics, Reports, Exports 4 views")
    add_bullet(doc, "Export Excel: Comprehensive, Executive, Employee, Self-Evaluation, Detailed Data, Individual, Weighted, Raw Data")
    add_bullet(doc, "Export PDF: Individual, Comprehensive")
    add_bullet(doc, "API endpoints: Dashboard data, Completion stats, Real-time data, Individual angle report")

    add_heading_styled(doc, "10.3 ระบบจัดการ Assignment (Assignment Management)", level=2)
    add_paragraph_styled(doc, "สถานะ: Completed", bold=True, color=RGBColor(0x00, 0x80, 0x00))
    add_bullet(doc, "Admin สามารถสร้าง, แก้ไข, ลบ evaluation assignments")
    add_bullet(doc, "Bulk operations (สร้าง/ลบหลายรายการ)")
    add_bullet(doc, "Analytics + Export")

    # Footer
    doc.add_paragraph()
    doc.add_paragraph()
    footer_line = doc.add_paragraph()
    footer_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_line.add_run("─" * 60)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("เอกสารนี้จัดทำเมื่อ 1 มีนาคม 2569 (ปรับปรุง 7 มีนาคม 2569)")
    run.font.name = 'TH Sarabun New'
    run.font.size = Pt(10)
    run.italic = True
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    footer2 = doc.add_paragraph()
    footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer2.add_run("ราคาอาจปรับเปลี่ยนได้หาก scope มีการเปลี่ยนแปลงหลังเริ่มงาน")
    run.font.name = 'TH Sarabun New'
    run.font.size = Pt(10)
    run.italic = True
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # Save
    output_path = os.path.join(os.path.dirname(__file__), 'project_proposal.docx')
    doc.save(output_path)
    print(f"Generated: {output_path}")
    return output_path

if __name__ == '__main__':
    generate_docx()
