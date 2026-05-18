#!/usr/bin/env python3
"""Generate evaluation forms analysis document (.docx)"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'TH Sarabun New'
font.size = Pt(14)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')

# ── Title ──
title = doc.add_heading('', level=0)
run = title.add_run('รายงานวิเคราะห์แบบประเมิน 360 องศา')
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0, 51, 102)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('การนิคมอุตสาหกรรมแห่งประเทศไทย (กนอ.)')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(80, 80, 80)

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub2.add_run('ปีงบประมาณ 2568 - 2569')
run2.font.size = Pt(14)
run2.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# Section 1: Overview
# ══════════════════════════════════════════════════════════════
doc.add_heading('1. ภาพรวมแบบประเมินทั้งหมด', level=1)

doc.add_paragraph(
    'ระบบประเมิน 360 องศา ของ กนอ. ประกอบด้วยแบบประเมิน 3 กลุ่มหลัก '
    'โดยแบ่งตามระดับตำแหน่ง (Grade) และประเภทผู้ประเมิน (ภายใน/ภายนอก/ประเมินตนเอง)'
)

# Summary table
tbl = doc.add_table(rows=1, cols=6)
tbl.style = 'Light Grid Accent 1'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = tbl.rows[0].cells
headers = ['Eval ID', 'ชื่อแบบประเมิน', 'ระดับ', 'ประเภท', 'จำนวนคำถาม', 'ปีงบ']
for i, h in enumerate(headers):
    hdr[i].text = h
    for p in hdr[i].paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.bold = True
            r.font.size = Pt(13)

evals_data = [
    # FY 2568
    ('1', 'ประเมินผู้บริหาร 9-12 (ภายใน)', '9-12', 'ภายใน', '39', '2568'),
    ('2', 'ประเมินผู้บริหาร 9-12 (ภายนอก)', '9-12', 'ภายนอก', '34', '2568'),
    ('3', 'ประเมินพนักงาน 5-8', '5-8', 'ภายใน', '29', '2568'),
    ('4', 'ประเมินตนเอง 9-12', '9-12', 'ตนเอง', '32', '2568'),
    ('5', 'ประเมินตนเอง 5-8', '5-8', 'ตนเอง', '22', '2568'),
    # FY 2569
    ('33', 'ประเมินผู้ว่าการ (ภายใน)', '13', 'ภายใน', '31', '2569'),
    ('34', 'ประเมินผู้ว่าการ (ภายนอก)', '13', 'ภายนอก', '28', '2569'),
    ('35', 'ประเมินผู้บริหาร 9-12 (ภายใน)', '9-12', 'ภายใน', '46', '2569'),
    ('36', 'ประเมินผู้บริหาร 9-12 (ภายนอก)', '9-12', 'ภายนอก', '42', '2569'),
    ('37', 'ประเมินพนักงาน 4-8', '4-8', 'ภายใน', '39', '2569'),
    ('38', 'ประเมินตนเอง ผู้ว่าการ', '13', 'ตนเอง', '31', '2569'),
    ('39', 'ประเมินตนเอง 9-12', '9-12', 'ตนเอง', '32', '2569'),
    ('40', 'ประเมินตนเอง 4-8', '4-8', 'ตนเอง', '22', '2569'),
]

for row_data in evals_data:
    row = tbl.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val
        for p in row[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.size = Pt(12)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# Section 2: Evaluation Selection Logic
# ══════════════════════════════════════════════════════════════
doc.add_heading('2. วิธีการเลือกแบบประเมิน', level=1)

doc.add_heading('2.1 หลักการเลือกแบบประเมิน', level=2)
doc.add_paragraph(
    'ระบบใช้ EvaluationLookupService เป็นศูนย์กลางในการเลือกแบบประเมิน '
    'โดยพิจารณาจาก 3 ปัจจัยหลัก:'
)

factors = [
    'ระดับตำแหน่ง (Grade) ของผู้ถูกประเมิน — ตรวจสอบจาก grade_min และ grade_max ในตาราง evaluations',
    'ประเภทผู้ประเมิน (user_type) — internal (บุคลากรภายใน) หรือ external (บุคลากรภายนอก)',
    'ประเภทการประเมิน — ประเมินตนเอง (title มีคำว่า "ประเมินตนเอง") หรือประเมินผู้อื่น',
]
for f in factors:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('2.2 การจับกลุ่มระดับตำแหน่ง (Grade Group)', level=2)

tbl2 = doc.add_table(rows=1, cols=4)
tbl2.style = 'Light Grid Accent 1'
hdr2 = tbl2.rows[0].cells
for i, h in enumerate(['กลุ่ม', 'ระดับ', 'องศาที่รองรับ', 'หมายเหตุ']):
    hdr2[i].text = h
    for p in hdr2[i].paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.size = Pt(13)

grade_groups = [
    ('ผู้ว่าการ', '13', 'self, top, bottom, left, right', 'มีแบบประเมินเฉพาะ (ปี 69+)'),
    ('ผู้บริหาร', '9-12', 'self, top, bottom, left, right', 'ทุกองศา'),
    ('พนักงาน', '4-8', 'self, top, left', 'ไม่มี bottom, right'),
]
for g in grade_groups:
    row = tbl2.add_row().cells
    for i, val in enumerate(g):
        row[i].text = val
        for p in row[i].paragraphs:
            for r in p.runs:
                r.font.size = Pt(12)

doc.add_paragraph()

doc.add_heading('2.3 ลำดับการเลือกแบบประเมิน (Priority)', level=2)

doc.add_heading('กรณีประเมินผู้อื่น (Assigned Evaluation)', level=3)
steps_assigned = [
    'ตรวจสอบ evaluation_id ที่บันทึกไว้ใน evaluation_assignments — ถ้ามีและยังเป็น published ใช้ eval นั้น',
    'ถ้าไม่มี: ใช้ EvaluationLookupService::findByGrade() — หาจาก grade ของผู้ถูกประเมิน + user_type',
    'เลือก evaluation ที่ status=published, grade_min <= grade <= grade_max, ไม่มีคำว่า "ประเมินตนเอง"',
    'ถ้ามีหลายแบบตรงเงื่อนไข: เลือกอันล่าสุด (latest)',
]
for i, s in enumerate(steps_assigned, 1):
    doc.add_paragraph(f'{i}. {s}')

doc.add_heading('กรณีประเมินตนเอง (Self Evaluation)', level=3)
steps_self = [
    'ใช้ EvaluationLookupService::findSelfEvalByGrade() — หาจาก grade ของผู้ใช้',
    'ลำดับแรก: หาแบบที่ title มีคำว่า "ประเมินตนเอง" + grade ตรง',
    'ถ้าไม่มี: Fallback ไปใช้แบบประเมินพนักงาน (internal) ปกติ',
]
for i, s in enumerate(steps_self, 1):
    doc.add_paragraph(f'{i}. {s}')

doc.add_heading('กรณี Admin สร้างคู่ประเมิน', level=3)
doc.add_paragraph(
    'เมื่อ Admin สร้าง assignment ใหม่ ระบบจะหา evaluation form ที่ตรง grade ของผู้ถูกประเมิน '
    'แล้วบันทึก evaluation_id ลงใน evaluation_assignments table โดยอัตโนมัติ'
)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# Section 3: FY 2568 Detail
# ══════════════════════════════════════════════════════════════
doc.add_heading('3. รายละเอียดแบบประเมิน ปีงบประมาณ 2568', level=1)

# ── Eval 1 ──
doc.add_heading('3.1 ประเมินผู้บริหาร 9-12 ภายใน (Eval 1) — 39 คำถาม', level=2)
doc.add_paragraph('ใช้สำหรับ: บุคลากรภายใน ประเมินผู้บริหารระดับ 9-12')

tbl_e1 = doc.add_table(rows=1, cols=4)
tbl_e1.style = 'Light List Accent 1'
for i, h in enumerate(['ส่วน', 'ด้าน', 'ประเภทคำถาม', 'จำนวน']):
    tbl_e1.rows[0].cells[i].text = h
    for r in tbl_e1.rows[0].cells[i].paragraphs[0].runs:
        r.font.bold = True; r.font.size = Pt(12)

rows_e1 = [
    ('ส่วนที่ 1\nการประเมิน 360 องศา\n(23 ข้อ)', 'ด้านความเป็นผู้นำ (Leadership)', 'Rating 1-5', '4'),
    ('', 'ด้านการมีวิสัยทัศน์ (Vision)', 'Rating 1-5', '3'),
    ('', 'ด้านการติดต่อสื่อสาร (Communication)', 'Rating 1-5', '4'),
    ('', 'ด้านความสามารถในการคิดและนวัตกรรม', 'Rating 1-5', '4'),
    ('', 'ด้านจริยธรรมในการปฏิบัติงาน', 'Rating 1-5', '4'),
    ('', 'ด้านทักษะระหว่างบุคคลและความร่วมมือ', 'Rating 1-5', '4'),
    ('ส่วนที่ 2\nวัฒนธรรมองค์กร\n(16 ข้อ)', '2.1 การยอมรับพฤติกรรมตามค่านิยม I-EA-T', 'Rating 1-5', '8'),
    ('', '2.2 การแสดงพฤติกรรมตามค่านิยม I-EA-T', 'Rating 1-5', '8'),
]
for rd in rows_e1:
    row = tbl_e1.add_row().cells
    for i, v in enumerate(rd):
        row[i].text = v
        for r in row[i].paragraphs[0].runs: r.font.size = Pt(11)

doc.add_paragraph()

# ── Eval 3 ──
doc.add_heading('3.2 ประเมินพนักงาน 5-8 ภายใน (Eval 3) — 29 คำถาม', level=2)
doc.add_paragraph('ใช้สำหรับ: บุคลากรภายใน ประเมินพนักงานระดับ 5-8')

tbl_e3 = doc.add_table(rows=1, cols=4)
tbl_e3.style = 'Light List Accent 1'
for i, h in enumerate(['ส่วน', 'ด้าน', 'ประเภทคำถาม', 'จำนวน']):
    tbl_e3.rows[0].cells[i].text = h
    for r in tbl_e3.rows[0].cells[i].paragraphs[0].runs:
        r.font.bold = True; r.font.size = Pt(12)

rows_e3 = [
    ('ส่วนที่ 1\nค่านิยมและสมรรถนะ\n(13 ข้อ)', 'ด้านเก่งคิด (IQ)', 'Rating 1-5', '3'),
    ('', 'ด้านเก่งคน (EQ)', 'Rating 1-5', '3'),
    ('', 'ด้านเก่งงาน (AQ + TQ)', 'Rating 1-5', '4'),
    ('', 'ด้านความยั่งยืน (Sustainability)', 'Rating 1-5', '3'),
    ('ส่วนที่ 2\nวัฒนธรรมองค์กร\n(16 ข้อ)', '2.1 การยอมรับพฤติกรรมตามค่านิยม I-EA-T', 'Rating 1-5', '8'),
    ('', '2.2 การแสดงพฤติกรรมตามค่านิยม I-EA-T', 'Rating 1-5', '8'),
]
for rd in rows_e3:
    row = tbl_e3.add_row().cells
    for i, v in enumerate(rd):
        row[i].text = v
        for r in row[i].paragraphs[0].runs: r.font.size = Pt(11)

doc.add_paragraph()

# ── Eval 4 ──
doc.add_heading('3.3 ประเมินตนเอง 9-12 (Eval 4) — 32 คำถาม', level=2)
doc.add_paragraph('ใช้สำหรับ: ผู้บริหาร 9-12 ประเมินตนเอง')

tbl_e4 = doc.add_table(rows=1, cols=4)
tbl_e4.style = 'Light List Accent 1'
for i, h in enumerate(['ส่วน', 'ด้าน', 'ประเภทคำถาม', 'จำนวน']):
    tbl_e4.rows[0].cells[i].text = h
    for r in tbl_e4.rows[0].cells[i].paragraphs[0].runs:
        r.font.bold = True; r.font.size = Pt(12)

rows_e4 = [
    ('ส่วนที่ 1\nการประเมิน 360 องศา\n(23 ข้อ)', 'ด้านความเป็นผู้นำ (Leadership)', 'Rating 1-5', '4'),
    ('', 'ด้านการมีวิสัยทัศน์ (Vision)', 'Rating 1-5', '3'),
    ('', 'ด้านการติดต่อสื่อสาร (Communication)', 'Rating 1-5', '4'),
    ('', 'ด้านความสามารถในการคิดและนวัตกรรม', 'Rating 1-5', '4'),
    ('', 'ด้านจริยธรรมในการปฏิบัติงาน', 'Rating 1-5', '4'),
    ('', 'ด้านทักษะระหว่างบุคคลและความร่วมมือ', 'Rating 1-5', '4'),
    ('ส่วนที่ 2\nวัฒนธรรมองค์กร\n(8 ข้อ)', '2.1 การรับรู้ค่านิยม I-EA-T', 'Choice / Multiple Choice', '4'),
    ('', '2.2 ความเข้าใจค่านิยม I-EA-T', 'Choice', '2'),
    ('', '2.3 ความเข้าใจพฤติกรรมพึงประสงค์', 'Choice', '2'),
    ('ส่วนที่ 3\nปลายเปิด (1 ข้อ)', 'จงยกตัวอย่างพฤติกรรมพึงประสงค์ AQ', 'Open Text', '1'),
]
for rd in rows_e4:
    row = tbl_e4.add_row().cells
    for i, v in enumerate(rd):
        row[i].text = v
        for r in row[i].paragraphs[0].runs: r.font.size = Pt(11)

doc.add_paragraph()

# ── Eval 5 ──
doc.add_heading('3.4 ประเมินตนเอง 5-8 (Eval 5) — 22 คำถาม', level=2)
doc.add_paragraph('ใช้สำหรับ: พนักงาน 5-8 ประเมินตนเอง — โครงสร้างเหมือน Eval 4 แต่ส่วนที่ 1 ใช้คำถามเดียวกับ Eval 3 (13 ข้อ)')

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# Section 4: FY 2569 Detail
# ══════════════════════════════════════════════════════════════
doc.add_heading('4. รายละเอียดแบบประเมิน ปีงบประมาณ 2569', level=1)

# ── Eval 33 ──
doc.add_heading('4.1 ประเมินผู้ว่าการ ภายใน (Eval 33) — 31 คำถาม', level=2)
doc.add_paragraph('ใช้สำหรับ: บุคลากรภายใน ประเมินผู้ว่าการ กนอ. (Grade 13) — แบบใหม่ปี 69')

tbl_e33 = doc.add_table(rows=1, cols=4)
tbl_e33.style = 'Light List Accent 1'
for i, h in enumerate(['ส่วน', 'ด้าน', 'ประเภทคำถาม', 'จำนวน']):
    tbl_e33.rows[0].cells[i].text = h
    for r in tbl_e33.rows[0].cells[i].paragraphs[0].runs:
        r.font.bold = True; r.font.size = Pt(12)

rows_e33 = [
    ('ส่วนที่ 1\nการประเมิน 360 องศา\n(28 ข้อ)', 'ด้านความเป็นผู้นำ/บริหารจัดการ', 'Rating 1-5', '6'),
    ('', 'ด้านวิสัยทัศน์และกลยุทธ์', 'Rating 1-5', '5'),
    ('', 'ด้านความสามารถในการสื่อสาร', 'Rating 1-5', '4'),
    ('', 'ด้าน Mindset / Creativity / Innovation', 'Rating 1-5', '6'),
    ('', 'ด้านจริยธรรมในการปฏิบัติงาน', 'Rating 1-5', '4'),
    ('', 'ด้านความสัมพันธ์และการทำงานร่วมกับผู้อื่น', 'Rating 1-5', '3'),
    ('ส่วนที่ 2\nปลายเปิด (3 ข้อ)', 'จุดแข็ง / สิ่งที่ควรพัฒนา / อื่นๆ', 'Open Text', '3'),
]
for rd in rows_e33:
    row = tbl_e33.add_row().cells
    for i, v in enumerate(rd):
        row[i].text = v
        for r in row[i].paragraphs[0].runs: r.font.size = Pt(11)

doc.add_paragraph()

# ── Eval 35 ──
doc.add_heading('4.2 ประเมินผู้บริหาร 9-12 ภายใน (Eval 35) — 46 คำถาม', level=2)
doc.add_paragraph('ใช้สำหรับ: บุคลากรภายใน ประเมินผู้บริหารระดับ 9-12 — เพิ่ม 2 ด้านใหม่จากปี 68')

tbl_e35 = doc.add_table(rows=1, cols=4)
tbl_e35.style = 'Light List Accent 1'
for i, h in enumerate(['ส่วน', 'ด้าน', 'ประเภทคำถาม', 'จำนวน']):
    tbl_e35.rows[0].cells[i].text = h
    for r in tbl_e35.rows[0].cells[i].paragraphs[0].runs:
        r.font.bold = True; r.font.size = Pt(12)

rows_e35 = [
    ('ส่วนที่ 1\nการประเมิน 360 องศา\n(30 ข้อ)', 'ด้านความเป็นผู้นำ (Leadership)', 'Rating 1-5', '5'),
    ('', 'ด้านการมีวิสัยทัศน์ (Vision)', 'Rating 1-5', '3'),
    ('', 'ด้านการติดต่อสื่อสาร (Communication)', 'Rating 1-5', '4'),
    ('', 'ด้านความสามารถในการคิดและนวัตกรรม', 'Rating 1-5', '6'),
    ('', 'ด้านจริยธรรมในการปฏิบัติงาน', 'Rating 1-5', '4'),
    ('', 'ด้านทักษะระหว่างบุคคลและความร่วมมือ', 'Rating 1-5', '4'),
    ('', '★ ด้านการบริหารความเสี่ยงและความรับผิดชอบ (ใหม่)', 'Rating 1-5', '4'),
    ('ส่วนที่ 2\nวัฒนธรรมองค์กร\n(16 ข้อ)', 'ระดับความสำคัญ', 'Rating 1-5', '8'),
    ('', 'ระดับการแสดงพฤติกรรม', 'Rating 1-5', '8'),
]
for rd in rows_e35:
    row = tbl_e35.add_row().cells
    for i, v in enumerate(rd):
        row[i].text = v
        for r in row[i].paragraphs[0].runs: r.font.size = Pt(11)

doc.add_paragraph()

# ── Eval 37 ──
doc.add_heading('4.3 ประเมินพนักงาน 4-8 (Eval 37) — 39 คำถาม', level=2)
doc.add_paragraph('ใช้สำหรับ: บุคลากรภายใน ประเมินพนักงานระดับ 4-8 — เพิ่ม 2 ด้านใหม่ + ขยาย grade จาก 5-8 เป็น 4-8')

tbl_e37 = doc.add_table(rows=1, cols=4)
tbl_e37.style = 'Light List Accent 1'
for i, h in enumerate(['ส่วน', 'ด้าน', 'ประเภทคำถาม', 'จำนวน']):
    tbl_e37.rows[0].cells[i].text = h
    for r in tbl_e37.rows[0].cells[i].paragraphs[0].runs:
        r.font.bold = True; r.font.size = Pt(12)

rows_e37 = [
    ('ส่วนที่ 1\nค่านิยมและสมรรถนะ\n(23 ข้อ)', 'ด้านเก่งคิด (IQ)', 'Rating 1-5', '4'),
    ('', 'ด้านเก่งคน (EQ)', 'Rating 1-5', '4'),
    ('', 'ด้านเก่งงาน (AQ + TQ)', 'Rating 1-5', '4'),
    ('', 'ด้านความยั่งยืน (Sustainability)', 'Rating 1-5', '3'),
    ('', '★ ด้านการคิดเชิงนวัตกรรมและการแก้ปัญหา (ใหม่)', 'Rating 1-5', '4'),
    ('', '★ ด้านการเรียนรู้และปรับตัวอย่างต่อเนื่อง (ใหม่)', 'Rating 1-5', '4'),
    ('ส่วนที่ 2\nวัฒนธรรมองค์กร\n(16 ข้อ)', 'ระดับความสำคัญ', 'Rating 1-5', '8'),
    ('', 'ระดับการแสดงพฤติกรรม', 'Rating 1-5', '8'),
]
for rd in rows_e37:
    row = tbl_e37.add_row().cells
    for i, v in enumerate(rd):
        row[i].text = v
        for r in row[i].paragraphs[0].runs: r.font.size = Pt(11)

doc.add_paragraph()

# ── Self-eval 38-40 ──
doc.add_heading('4.4 ประเมินตนเอง (Eval 38, 39, 40)', level=2)
doc.add_paragraph(
    'แบบประเมินตนเองทั้ง 3 ระดับ (ผู้ว่าการ, 9-12, 4-8) มีโครงสร้างเหมือนกัน:'
)

tbl_self = doc.add_table(rows=1, cols=3)
tbl_self.style = 'Light List Accent 1'
for i, h in enumerate(['ส่วน', 'เนื้อหา', 'ประเภทคำถาม']):
    tbl_self.rows[0].cells[i].text = h
    for r in tbl_self.rows[0].cells[i].paragraphs[0].runs:
        r.font.bold = True; r.font.size = Pt(12)

self_rows = [
    ('ส่วนที่ 1: Rating', 'คำถาม Rating เดียวกับแบบประเมินพนักงาน\n(แต่ใช้สรรพนาม "ท่าน" สำหรับผู้ว่าการ)', 'Rating 1-5'),
    ('ส่วนที่ 2: วัฒนธรรมองค์กร (8 ข้อ)', '2.1 การรับรู้ค่านิยม (4 ข้อ Choice/MC)\n2.2 ความเข้าใจค่านิยม (2 ข้อ Choice)\n2.3 พฤติกรรมพึงประสงค์ (2 ข้อ Choice)', 'Choice / Multiple Choice'),
    ('ส่วนที่ 3: ปลายเปิด (1 ข้อ)', 'จงยกตัวอย่างพฤติกรรมพึงประสงค์\nของค่านิยม "เก่งงาน AQ"', 'Open Text'),
]
for rd in self_rows:
    row = tbl_self.add_row().cells
    for i, v in enumerate(rd):
        row[i].text = v
        for r in row[i].paragraphs[0].runs: r.font.size = Pt(11)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# Section 5: Key Differences
# ══════════════════════════════════════════════════════════════
doc.add_heading('5. ความแตกต่างระหว่างแบบประเมินพนักงาน vs ประเมินตนเอง', level=1)

tbl_diff = doc.add_table(rows=1, cols=3)
tbl_diff.style = 'Light Grid Accent 1'
for i, h in enumerate(['หัวข้อ', 'แบบประเมินพนักงาน', 'แบบประเมินตนเอง']):
    tbl_diff.rows[0].cells[i].text = h
    for r in tbl_diff.rows[0].cells[i].paragraphs[0].runs:
        r.font.bold = True; r.font.size = Pt(13)

diff_rows = [
    ('ส่วนที่ 1 (Rating)', 'คำถาม Rating เหมือนกัน\n(ใช้สรรพนาม "ผู้บริหาร/พนักงาน")', 'คำถาม Rating เหมือนกัน\n(ใช้สรรพนาม "ท่าน" สำหรับผู้ว่าการ)'),
    ('ส่วนที่ 2 (วัฒนธรรมองค์กร)', '16 ข้อ Rating\n• 8 ข้อ ระดับยอมรับ/ความสำคัญ\n• 8 ข้อ ระดับแสดงพฤติกรรม', '8 ข้อ Choice/MC\n• 4 ข้อ การรับรู้ค่านิยม\n• 2 ข้อ ความเข้าใจค่านิยม\n• 2 ข้อ พฤติกรรมพึงประสงค์'),
    ('ส่วนที่ 3 (ปลายเปิด)', 'ไม่มี\n(ยกเว้นผู้ว่าการ: 3 ข้อปลายเปิด)', 'มี 1 ข้อ\n(ยกตัวอย่างพฤติกรรม AQ)'),
    ('จำนวนคำถามรวม (9-12)', '39 ข้อ (ปี 68)\n46 ข้อ (ปี 69)', '32 ข้อ (ทั้ง 2 ปี)'),
    ('จำนวนคำถามรวม (4-8)', '29 ข้อ (ปี 68)\n39 ข้อ (ปี 69)', '22 ข้อ (ทั้ง 2 ปี)'),
]
for rd in diff_rows:
    row = tbl_diff.add_row().cells
    for i, v in enumerate(rd):
        row[i].text = v
        for r in row[i].paragraphs[0].runs: r.font.size = Pt(11)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# Section 6: Changes from 68 to 69
# ══════════════════════════════════════════════════════════════
doc.add_heading('6. การเปลี่ยนแปลงจากปี 2568 สู่ 2569', level=1)

changes = [
    'เพิ่มแบบประเมินผู้ว่าการ กนอ. แยกจากกลุ่ม 9-12 (Eval 33, 34, 38)',
    'ขยายช่วง Grade จาก 5-8 เป็น 4-8 — รองรับพนักงานระดับ 4',
    'ผู้บริหาร 9-12: เพิ่มด้าน "การบริหารความเสี่ยงและความรับผิดชอบ" (4 ข้อ) และเพิ่มข้อในด้านอื่นๆ รวม 30→46 ข้อ',
    'พนักงาน 4-8: เพิ่มด้าน "การคิดเชิงนวัตกรรม" และ "การเรียนรู้และปรับตัว" (8 ข้อ) รวม 29→39 ข้อ',
    'ผู้ว่าการภายใน: ใช้คำถามเฉพาะ (ผวก. กนอ.) + 3 ข้อปลายเปิด ไม่มีส่วนวัฒนธรรมองค์กร',
    'ประเมินตนเองทุกระดับ: ส่วนที่ 2 ยังคงเป็น 8 ข้อ Choice/MC เหมือนปี 68',
]
for c in changes:
    doc.add_paragraph(c, style='List Bullet')

# ── Save ──
output_path = os.path.join(os.path.dirname(__file__), 'evaluation_forms_analysis.docx')
doc.save(output_path)
print(f'Saved to: {output_path}')
