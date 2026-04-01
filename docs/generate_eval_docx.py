#!/usr/bin/env python3
"""Generate .docx files for all year 69 evaluations from DB JSON."""

import json
import os
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
DB_JSON = os.path.join(SCRIPT_DIR, 'db_evals_final.json')
OUTPUT_DIR = os.path.join(SCRIPT_DIR, 'แบบประเมินจากระบบ_ปี69')


def set_cell_shading(cell, color):
    """Set cell background color."""
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear',
    })
    shading_elm.append(shading)


def add_rating_table(doc, questions, aspect_groups, header_label="ระดับความคิดเห็น"):
    """Add a rating table (5-point scale) for Part 1 style questions."""
    # Create table: question + 5 rating columns
    num_rows = 2  # header rows
    for aspect_name, qs in aspect_groups:
        num_rows += 1  # aspect header
        num_rows += len(qs)  # questions

    table = doc.add_table(rows=2, cols=6)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Cm(12)
        for i in range(1, 6):
            row.cells[i].width = Cm(1.2)

    # Header row 1
    h1 = table.rows[0]
    h1.cells[0].text = "หัวข้อและประเด็นการประเมิน ประจำปี 2569"
    for i in range(1, 6):
        h1.cells[i].text = header_label
    # Merge rating header
    h1.cells[1].merge(h1.cells[5])

    # Header row 2
    h2 = table.rows[1]
    h2.cells[0].text = "หัวข้อและประเด็นการประเมิน ประจำปี 2569"
    for i, score in enumerate([5, 4, 3, 2, 1]):
        h2.cells[i + 1].text = str(score)

    # Style headers
    for row in [h1, h2]:
        for cell in row.cells:
            set_cell_shading(cell, "4472C4")
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.size = Pt(10)

    # Add aspect groups and questions
    q_num = 1
    aspect_num = 1
    for aspect_name, qs in aspect_groups:
        # Aspect header row
        row = table.add_row()
        row.cells[0].text = f"{aspect_num}) {aspect_name}"
        row.cells[0].merge(row.cells[5])
        set_cell_shading(row.cells[0], "D6E4F0")
        for p in row.cells[0].paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)

        # Question rows
        for q in qs:
            row = table.add_row()
            row.cells[0].text = f"{q_num}. {q['title']}"
            for p in row.cells[0].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
            for i in range(1, 6):
                row.cells[i].text = ""
                for p in row.cells[i].paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            q_num += 1
        aspect_num += 1

    return table


def add_culture_rating_table(doc, questions, aspect_name, header_label):
    """Add a culture assessment rating table."""
    table = doc.add_table(rows=2, cols=6)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    h1 = table.rows[0]
    h1.cells[0].text = "ประเด็นการประเมิน ประจำปี 2569"
    for i in range(1, 6):
        h1.cells[i].text = header_label
    h1.cells[1].merge(h1.cells[5])

    h2 = table.rows[1]
    h2.cells[0].text = "ประเด็นการประเมิน ประจำปี 2569"
    for i, score in enumerate([5, 4, 3, 2, 1]):
        h2.cells[i + 1].text = str(score)

    for row in [h1, h2]:
        for cell in row.cells:
            set_cell_shading(cell, "4472C4")
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.size = Pt(10)

    for i, q in enumerate(questions):
        row = table.add_row()
        row.cells[0].text = f"{i + 1}. {q['title']}"
        for p in row.cells[0].paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)

    return table


def add_choice_section(doc, aspect_name, questions):
    """Add choice/multiple choice questions as paragraphs."""
    doc.add_paragraph(aspect_name, style='Heading 3')

    for i, q in enumerate(questions):
        q_type = q['type']
        title = q['title']

        p = doc.add_paragraph()
        run = p.add_run(f"{i + 1}) {title}")
        run.font.size = Pt(10)

        if q_type in ('choice', 'multiple_choice'):
            note = "(เลือกได้มากกว่า 1 ข้อ)" if q_type == 'multiple_choice' else ""
            if note:
                run2 = p.add_run(f" {note}")
                run2.font.size = Pt(9)
                run2.font.italic = True

            # Add placeholder choices
            for opt_num in range(1, 4):
                opt_p = doc.add_paragraph(style='List Number')
                opt_p.text = f"    ({opt_num}) ตัวเลือก {opt_num}"
                for run in opt_p.runs:
                    run.font.size = Pt(10)

        elif q_type == 'open_text':
            doc.add_paragraph("    ____________________________________________________________")


def generate_eval_docx(eval_data, output_path):
    """Generate a .docx file for one evaluation."""
    doc = Document()

    # Title
    title_para = doc.add_heading(eval_data['title'], level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Info
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(
        f"user_type: {eval_data.get('user_type', 'N/A')} | "
        f"grade: {eval_data.get('grade_min', '?')}-{eval_data.get('grade_max', '?')} | "
        f"Eval ID: {eval_data['id']}"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)

    # Score description
    doc.add_paragraph()
    desc = doc.add_paragraph()
    desc.add_run("คำอธิบายค่าคะแนน:").bold = True
    scores = [
        ("5", "ดีเยี่ยม", "แสดงพฤติกรรมที่สอดคล้องกับ Core Values ในการเป็นต้นแบบได้อย่างโดดเด่น"),
        ("4", "ดีมาก", "แสดงพฤติกรรมในการดำเนินการประเด็นดังกล่าวเหนือกว่าความต้องการ"),
        ("3", "ดี", "แสดงพฤติกรรมที่สอดคล้องในระดับพื้นฐาน"),
        ("2", "ต้องปรับปรุง", "ควรพิจารณาเพื่อดำเนินการปรับปรุง/พัฒนา"),
        ("1", "ต้องปรับปรุงอย่างมาก", ""),
    ]
    for score, label, detail in scores:
        p = doc.add_paragraph()
        p.add_run(f"ระดับคะแนน {score} = {label}").bold = True
        if detail:
            p.add_run(f" ({detail})")
        for run in p.runs:
            run.font.size = Pt(10)

    # Process each part
    for part in eval_data['parts']:
        doc.add_page_break()
        doc.add_heading(part['title'], level=2)

        # Determine part type
        part_order = part['order']
        aspects = part['aspects']

        if not aspects:
            continue

        # Check what types of questions this part has
        all_types = set()
        for a in aspects:
            for q in a['questions']:
                all_types.add(q['type'])

        if 'rating' in all_types and part_order == 1:
            # Part 1: Rating table with aspect groups
            aspect_groups = []
            for a in aspects:
                rating_qs = [q for q in a['questions'] if q['type'] == 'rating']
                if rating_qs:
                    aspect_groups.append((a['name'], rating_qs))

            if aspect_groups:
                add_rating_table(doc, None, aspect_groups)

        elif 'rating' in all_types and part_order == 2:
            # Part 2 with rating: culture assessment tables
            for a in aspects:
                rating_qs = [q for q in a['questions'] if q['type'] == 'rating']
                if rating_qs:
                    add_culture_rating_table(doc, rating_qs, a['name'], a['name'])
                    doc.add_paragraph()

        elif 'choice' in all_types or 'multiple_choice' in all_types:
            # Part 2/3 with choice questions (self-eval)
            for a in aspects:
                add_choice_section(doc, a['name'], a['questions'])

        elif 'open_text' in all_types:
            # Open text part
            for a in aspects:
                doc.add_heading(a['name'], level=3)
                for q in a['questions']:
                    p = doc.add_paragraph()
                    p.add_run(q['title']).font.size = Pt(10)
                    doc.add_paragraph("____________________________________________________________")
                    doc.add_paragraph("____________________________________________________________")
                    doc.add_paragraph("____________________________________________________________")

        else:
            # Mixed or unknown - show all
            for a in aspects:
                if a['questions']:
                    # Check if has rating
                    rating_qs = [q for q in a['questions'] if q['type'] == 'rating']
                    choice_qs = [q for q in a['questions'] if q['type'] in ('choice', 'multiple_choice')]
                    text_qs = [q for q in a['questions'] if q['type'] == 'open_text']

                    if rating_qs:
                        add_culture_rating_table(doc, rating_qs, a['name'], a['name'])
                        doc.add_paragraph()

                    if choice_qs or text_qs:
                        add_choice_section(doc, a['name'], choice_qs + text_qs)

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("--- สร้างจากฐานข้อมูลระบบประเมิน 360 องศา (Production) วันที่ 27 มีนาคม 2569 ---")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(160, 160, 160)

    doc.save(output_path)
    return output_path


def main():
    with open(DB_JSON, 'r', encoding='utf-8') as f:
        db_evals = json.load(f)

    os.makedirs(OUTPUT_DIR, exist_ok=True)

    # Also need full eval data with user_type/grade from production
    # The final JSON may not have these - let's check
    for eval_data in db_evals:
        eid = eval_data['id']
        title = eval_data['title']

        # Add grade info based on known mapping
        grade_map = {
            35: ('internal', 9, 12),
            36: ('external', 9, 12),
            37: ('internal', 4, 8),
            38: ('internal', 13, 13),
            39: ('internal', 9, 12),
            40: ('internal', 4, 8),
        }
        if eid in grade_map:
            eval_data['user_type'] = grade_map[eid][0]
            eval_data['grade_min'] = grade_map[eid][1]
            eval_data['grade_max'] = grade_map[eid][2]

        # Clean filename
        safe_title = re.sub(r'[<>:"/\\|?*]', '', title)
        safe_title = safe_title[:80]
        filename = f"Eval{eid}_{safe_title}.docx"
        output_path = os.path.join(OUTPUT_DIR, filename)

        generate_eval_docx(eval_data, output_path)

        total_qs = sum(
            len(q['questions'])
            for p in eval_data['parts']
            for q in p['aspects']
        )
        print(f"✅ Eval {eid}: {title[:50]} ({total_qs} Qs) -> {filename}")

    print(f"\nAll files saved to: {OUTPUT_DIR}")


if __name__ == '__main__':
    main()
