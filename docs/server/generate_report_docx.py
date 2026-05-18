"""
Generate Load Test Report as .docx
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_table(doc, headers, rows, col_widths=None, header_color='4B0082'):
    """Add a formatted table"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_color)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            if c_idx == 0:
                run.bold = True
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'F3F0FF')

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table

def main():
    doc = Document()

    # Default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'TH Sarabun New'
    font.size = Pt(14)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')

    # ========== COVER ==========
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('รายงานผลทดสอบประสิทธิภาพระบบ')
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(75, 0, 130)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('(Load Test Report)')
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ระบบประเมินผล 360 องศา')
    run.font.size = Pt(18)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('สำหรับผู้ว่าการ กลุ่มผู้บริหารและพนักงาน ระดับ 4-12')
    run.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('การนิคมอุตสาหกรรมแห่งประเทศไทย (กนอ.)')
    run.font.size = Pt(16)
    run.bold = True

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('จัดทำโดย Miles Consult Group')
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('วันที่ 22 มีนาคม 2569')
    run.font.size = Pt(14)

    doc.add_page_break()

    # ========== SECTION 1 ==========
    h = doc.add_heading('1. วัตถุประสงค์', level=1)
    h.runs[0].font.color.rgb = RGBColor(75, 0, 130)

    doc.add_paragraph(
        'ทดสอบความสามารถในการรองรับผู้ใช้งานพร้อมกันของ Server ปัจจุบัน '
        'เพื่อประเมินว่าสามารถรองรับการใช้งานจริงในวันทดสอบระบบ '
        '(600 คนพร้อมกัน) ได้หรือไม่'
    )

    add_table(doc,
        ['รายการ', 'รายละเอียด'],
        [
            ['เป้าหมายทดสอบ', 'https://evaluation.milesconsult.com'],
            ['เครื่องมือ', 'Grafana k6 v0.54.0'],
            ['วันที่ทดสอบ', '22 มีนาคม 2569'],
            ['ทดสอบโดย', 'Miles Consult Group'],
        ],
        col_widths=[5, 12]
    )

    # ========== SECTION 2 ==========
    h = doc.add_heading('2. Spec Server ปัจจุบัน (Hostinger Cloud Enterprise)', level=1)
    h.runs[0].font.color.rgb = RGBColor(75, 0, 130)

    add_table(doc,
        ['รายการ', 'ค่า'],
        [
            ['Hosting', 'Hostinger Cloud Enterprise'],
            ['CPU', '6 Cores (Shared)'],
            ['RAM', '16 GB'],
            ['Storage', '200 GB NVMe SSD'],
            ['Bandwidth', 'ไม่จำกัด'],
            ['ราคา', '~$29.99/เดือน (~1,050 บาท)'],
        ],
        col_widths=[5, 12]
    )

    # ========== SECTION 3 ==========
    h = doc.add_heading('3. วิธีการทดสอบ', level=1)
    h.runs[0].font.color.rgb = RGBColor(75, 0, 130)

    doc.add_heading('Scenario 1: Capacity Test (หา Threshold)', level=2)
    doc.add_paragraph(
        'ค่อยๆ เพิ่มจำนวนผู้ใช้พร้อมกันจาก 10 → 25 → 50 → 75 → 100 → 150 → 200 → 300 → 500 คน '
        'โดยจำลองการเข้าหน้าเว็บ (หน้าแรก, หน้าเข้าสู่ระบบ, หน้าผู้ประเมินภายนอก) '
        'พร้อมจับเวลาตอบสนองและอัตรา error'
    )

    doc.add_heading('Scenario 2: Login Load Test (1,000 คนพร้อมกัน)', level=2)
    doc.add_paragraph(
        'จำลองผู้ใช้ 1,000 คนพร้อมกัน ทำการเข้าสู่ระบบ (GET login page → POST login) '
        'เพื่อทดสอบขีดจำกัดสูงสุด'
    )

    # ========== SECTION 4 ==========
    h = doc.add_heading('4. ผลการทดสอบ', level=1)
    h.runs[0].font.color.rgb = RGBColor(75, 0, 130)

    doc.add_heading('4.1 Capacity Test — ผลรายระดับ', level=2)

    add_table(doc,
        ['ผู้ใช้พร้อมกัน (VUs)', 'Throughput (req/10s)', 'Response Time เฉลี่ย', 'สถานะ'],
        [
            ['10', '6-7', '~64ms', '✅ ปกติ'],
            ['25', '10-12', '~64ms', '✅ ปกติ'],
            ['50', '22-30', '~64ms', '✅ ปกติ'],
            ['75', '33-41', '~64ms', '✅ ปกติ'],
            ['100', '43-51', '~64ms', '✅ ปกติ'],
            ['150', '66-78', '~80ms', '✅ ยังใช้ได้'],
            ['200', '88-98', '~120ms', '⚠️ เริ่มช้าลง'],
            ['300', '135-160', '~250ms+', '⚠️ Error เริ่มเพิ่ม'],
            ['500', '235-274', 'ไม่ตอบสนอง', '❌ Server ถูกบล็อก'],
        ],
        col_widths=[4.5, 4.5, 4.5, 4.5]
    )

    doc.add_heading('4.2 สรุปตัวเลขสำคัญ', level=2)

    add_table(doc,
        ['เกณฑ์', 'ค่าที่วัดได้'],
        [
            ['Response Time เฉลี่ย (ปกติ)', '64 ms'],
            ['Response Time p95', '100 ms'],
            ['Response Time p99', '592 ms'],
            ['RPS สูงสุดที่รองรับได้', '~73 requests/sec'],
            ['ผู้ใช้พร้อมกันที่ใช้งานได้สบาย', '50-75 คน'],
            ['ผู้ใช้พร้อมกันสูงสุดก่อน degradation', '100-150 คน'],
            ['ผู้ใช้พร้อมกันที่ระบบล่ม', '200+ คน'],
        ],
        col_widths=[8, 8]
    )

    doc.add_heading('4.3 Login Load Test — 1,000 คนพร้อมกัน', level=2)

    add_table(doc,
        ['เกณฑ์', 'ผล'],
        [
            ['Requests ทั้งหมด', '17,066'],
            ['สำเร็จ', '0 (0%)'],
            ['ล้มเหลว', '17,066 (100%)'],
            ['สาเหตุ', 'TLS Internal Error — Hostinger WAF/DDoS Protection บล็อก'],
        ],
        col_widths=[6, 11]
    )

    # ========== SECTION 5 ==========
    h = doc.add_heading('5. สาเหตุของข้อจำกัด', level=1)
    h.runs[0].font.color.rgb = RGBColor(75, 0, 130)

    p = doc.add_paragraph()
    run = p.add_run('ปัญหาหลักไม่ใช่ Performance ของ Application ')
    run.bold = True
    p.add_run('(ซึ่งตอบสนองเร็วมากที่ 64ms) แต่เป็น:')

    items = [
        ('Hostinger WAF/DDoS Protection', 'ระบบป้องกันของ Hostinger ตัด TLS connection เมื่อมี concurrent connections สูงเกินกำหนด (~150-200 connections)'),
        ('Shared Infrastructure', 'Cloud Enterprise ใช้ทรัพยากรร่วมกับลูกค้าอื่น ไม่สามารถ customize ได้'),
        ('ไม่มี Root Access', 'ไม่สามารถปรับแต่ง PHP-FPM workers, Nginx config, หรือเพิ่ม Redis cache ได้'),
    ]
    for title, desc in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(title + ' — ')
        run.bold = True
        p.add_run(desc)

    # ========== SECTION 6 ==========
    h = doc.add_heading('6. ผลกระทบต่อการใช้งานจริง', level=1)
    h.runs[0].font.color.rgb = RGBColor(75, 0, 130)

    p = doc.add_paragraph()
    run = p.add_run('กรณี: วันทดสอบระบบ 600 คนพร้อมกัน')
    run.bold = True
    run.font.size = Pt(14)

    add_table(doc,
        ['สถานการณ์', 'Server รับได้', 'ผล'],
        [
            ['600 คนเข้าพร้อมกันทันที', '~150 คน', '❌ ไม่ผ่าน — 450 คนเข้าไม่ได้'],
            ['600 คนกระจาย 30 นาที', '~20-30 concurrent', '✅ รับได้'],
            ['600 คนกระจาย 1 ชั่วโมง', '~10-15 concurrent', '✅ รับได้สบาย'],
        ],
        col_widths=[6, 4, 7]
    )

    p = doc.add_paragraph()
    run = p.add_run('สรุป: Server ปัจจุบันไม่สามารถรองรับ 600 คนพร้อมกันในสถานการณ์ peak ได้')
    run.bold = True
    run.font.color.rgb = RGBColor(180, 0, 0)

    # ========== SECTION 7 ==========
    h = doc.add_heading('7. ข้อเสนอแนะ: ย้ายไป VPS KVM4 ของ Hostinger', level=1)
    h.runs[0].font.color.rgb = RGBColor(75, 0, 130)

    doc.add_heading('7.1 เปรียบเทียบ Spec และราคา', level=2)

    add_table(doc,
        ['', 'Cloud Enterprise\n(ปัจจุบัน)', 'VPS KVM4\n(แนะนำ)', 'VPS KVM8'],
        [
            ['CPU', '6 Cores (Shared)', '4 Cores (Dedicated)', '8 Cores (Dedicated)'],
            ['RAM', '16 GB', '16 GB', '32 GB'],
            ['Storage', '200 GB NVMe', '200 GB NVMe', '400 GB NVMe'],
            ['Bandwidth', 'ไม่จำกัด*', '16 TB/เดือน', '32 TB/เดือน'],
            ['WAF Rate Limit', 'มี (จำกัด concurrent)', 'ไม่มี (ควบคุมเอง)', 'ไม่มี (ควบคุมเอง)'],
            ['Root Access', 'ไม่มี', 'มี', 'มี'],
            ['Custom Config', 'ไม่ได้', 'ได้เต็มที่', 'ได้เต็มที่'],
            ['ราคา/เดือน', '~1,050 บาท', '~1,085 บาท', '~1,890 บาท'],
            ['ราคา promo ปีแรก', '-', '~350 บาท/เดือน', '~1,400 บาท/เดือน'],
            ['รองรับ 600 concurrent', '❌ ไม่ได้', '✅ ได้', '✅ สบาย'],
        ],
        col_widths=[4, 4.5, 4.5, 4.5]
    )

    doc.add_heading('7.2 เหตุผลที่แนะนำ VPS KVM4', level=2)

    reasons = [
        ('ไม่มี WAF Rate Limit', 'สามารถรับ 600+ concurrent connections ได้จริง ไม่ถูกบล็อก'),
        ('Root Access', 'ปรับแต่ง PHP-FPM (เพิ่ม workers), Nginx (tuning), OPcache, Redis ได้เต็มที่'),
        ('ราคาใกล้เคียงกัน', '~$31/เดือน vs ~$30/เดือน (ต่างกันแค่ ~$1)'),
        ('Dedicated Resources', 'CPU ไม่ต้องแชร์กับลูกค้าอื่น ทำให้ performance คงที่'),
        ('Scalable', 'อัพเกรดเป็น KVM8 ได้ง่ายหากต้องการในอนาคต'),
    ]
    for i, (title, desc) in enumerate(reasons, 1):
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(title + ' — ')
        run.bold = True
        p.add_run(desc)

    doc.add_heading('7.3 ประมาณการรองรับ Concurrent Users หลังย้าย', level=2)

    add_table(doc,
        ['การปรับแต่ง', 'Concurrent Users ที่คาดว่ารับได้'],
        [
            ['VPS KVM4 + Default Config', '~300-400 คน'],
            ['VPS KVM4 + PHP-FPM Tuning + OPcache', '~500-700 คน'],
            ['VPS KVM4 + Redis Session/Cache', '~700-1,000 คน'],
            ['VPS KVM8 + Full Optimization', '~1,500+ คน'],
        ],
        col_widths=[9, 8]
    )

    # ========== SECTION 8 ==========
    h = doc.add_heading('8. แผนการย้าย Server', level=1)
    h.runs[0].font.color.rgb = RGBColor(75, 0, 130)

    add_table(doc,
        ['ขั้นตอน', 'รายละเอียด', 'ระยะเวลา'],
        [
            ['1', 'สั่งซื้อ VPS KVM4 และ Setup Server', '2-3 ชั่วโมง'],
            ['2', 'ติดตั้ง Nginx + PHP 8.2 + MariaDB + SSL', '1-2 ชั่วโมง'],
            ['3', 'ย้ายข้อมูลและ Code จาก Cloud Enterprise', '1 ชั่วโมง'],
            ['4', 'ทดสอบระบบบน VPS ใหม่', '1 ชั่วโมง'],
            ['5', 'เปลี่ยน DNS ชี้ไปที่ VPS ใหม่', '15 นาที (รอ propagate 1-24 ชม.)'],
            ['6', 'Load Test ยืนยันผล', '30 นาที'],
            ['', 'รวม', '~1 วันทำการ'],
        ],
        col_widths=[2, 10, 5]
    )

    # ========== SECTION 9 ==========
    h = doc.add_heading('9. สรุปและข้อเสนอ', level=1)
    h.runs[0].font.color.rgb = RGBColor(75, 0, 130)

    add_table(doc,
        ['รายการ', 'รายละเอียด'],
        [
            ['ปัญหา', 'Server ปัจจุบันรับ 600 คนพร้อมกันไม่ได้ เพราะ Hostinger WAF บล็อก'],
            ['แนวทาง', 'ย้ายไป Hostinger VPS KVM4'],
            ['ราคา', 'ใกล้เคียงกัน (~$31 vs ~$30/เดือน)'],
            ['ผลที่คาดว่าจะได้', 'รองรับ 600-1,000 คนพร้อมกัน หลังปรับแต่ง'],
            ['ระยะเวลาย้าย', '1 วันทำการ'],
            ['ความเสี่ยง', 'ต่ำ — ย้ายภายใน Hostinger เดียวกัน, มี backup ตลอด'],
        ],
        col_widths=[5, 12]
    )

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('— จัดทำโดย Miles Consult Group —')
    run.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)

    # Save
    output = 'docs/server/load-test-report.docx'
    doc.save(output)
    print(f'Generated: {output}')

if __name__ == '__main__':
    main()
