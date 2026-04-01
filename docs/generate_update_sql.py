#!/usr/bin/env python3
"""Generate SQL to update evaluation questions to match Word documents."""

import os
import re
import json
from docx import Document

FOLDER = os.path.join(os.path.dirname(__file__), 'แบบประเมินปรับปรุง')
DB_JSON = os.path.join(os.path.dirname(__file__), 'db_evals_prod.json')


def normalize(text):
    return re.sub(r'\s+', ' ', text).strip()


def escape_sql(text):
    return text.replace("\\", "\\\\").replace("'", "\\'")


def get_db_eval(db_evals, eval_id):
    for e in db_evals:
        if e['id'] == eval_id:
            return e
    return None


def extract_part1_questions(filepath):
    doc = Document(filepath)
    questions = []
    current_aspect = None
    if not doc.tables:
        return questions
    table = doc.tables[0]
    for row in table.rows:
        cells = [normalize(cell.text) for cell in row.cells]
        text = cells[0]
        if not text or 'ระดับ' in text or 'หัวข้อและประเด็น' in text:
            continue
        if re.match(r'^\d+\)', text):
            current_aspect = re.sub(r'^\d+\)\s*', '', text).strip()
            continue
        m = re.match(r'^(\d+)\.\s*(.+)', text, re.DOTALL)
        if m:
            q_text = normalize(m.group(2))
            questions.append({
                'num': int(m.group(1)),
                'aspect': current_aspect,
                'text': q_text
            })
    return questions


def get_db_part1_rating_questions(eval_data):
    questions = []
    for part in eval_data['parts']:
        if part['order'] == 1:
            for aspect in part['aspects']:
                for q in aspect['questions']:
                    if q['type'] == 'rating':
                        questions.append({
                            'id': q['id'],
                            'aspect_name': aspect['name'],
                            'aspect_id': aspect['id'],
                            'text': normalize(q['title']),
                        })
    return questions


def main():
    with open(DB_JSON, 'r', encoding='utf-8') as f:
        db_evals = json.load(f)

    sql = []
    sql.append("-- ======================================================")
    sql.append("-- UPDATE EVALUATION QUESTIONS TO MATCH WORD DOCS (ปี 69)")
    sql.append("-- Production: u917560495_milesconsultdb")
    sql.append("-- Generated: 2026-03-27")
    sql.append("-- ======================================================")
    sql.append("")
    sql.append("-- STEP 0: Backup questions")
    sql.append("CREATE TABLE IF NOT EXISTS questions_backup_20260327 AS SELECT * FROM questions WHERE id BETWEEN 1322 AND 1600;")
    sql.append("CREATE TABLE IF NOT EXISTS aspects_backup_20260327 AS SELECT * FROM aspects WHERE id BETWEEN 270 AND 330;")
    sql.append("")

    # ---- Load Word files ----
    word_int912 = extract_part1_questions(
        os.path.join(FOLDER, 'แบบประเมิน 360 องศา สำหรับกลุ่มผู้บริหารระดับ 9-12 สำหรับบุคลากรภายใน.docx'))
    word_ext912 = extract_part1_questions(
        os.path.join(FOLDER, 'แบบประเมิน 360 องศา สำหรับกลุ่มผู้บริหารระดับ 9-12 สำหรับบุคลากรภายนอก.docx'))
    word_int48 = extract_part1_questions(
        os.path.join(FOLDER, 'แบบประเมิน 360 องศา สำหรับพนักงานระดับ 4-8 สำหรับพนักงาน.docx'))
    word_gov = extract_part1_questions(
        os.path.join(FOLDER, 'แบบประเมิน 360 องศา สำหรับผู้ว่าการ กนอ. (บุคลากรภายใน-ภายนอก).docx'))

    print(f"Word Int 9-12: {len(word_int912)} Qs")
    print(f"Word Ext 9-12: {len(word_ext912)} Qs")
    print(f"Word Int 4-8:  {len(word_int48)} Qs")
    print(f"Word Governor: {len(word_gov)} Qs")

    # ===========================
    # EVAL 35: Internal 9-12
    # ===========================
    db35 = get_db_part1_rating_questions(get_db_eval(db_evals, 35))
    sql.append("-- ============================================")
    sql.append("-- EVAL 35: Internal 9-12 (30 Qs, fix Q7)")
    sql.append("-- ============================================")

    for i, wq in enumerate(word_int912):
        if i < len(db35):
            w_text = normalize(wq['text'])
            d_text = normalize(db35[i]['text'])
            if w_text != d_text:
                sql.append(f"UPDATE questions SET title = '{escape_sql(w_text)}' WHERE id = {db35[i]['id']};")
    sql.append("")

    # ===========================
    # EVAL 36: External 9-12 (25 Qs in Word, 26 in DB - delete Q1378)
    # ===========================
    db36 = get_db_part1_rating_questions(get_db_eval(db_evals, 36))
    sql.append("-- ============================================")
    sql.append("-- EVAL 36: External 9-12 (delete 1 Q + update)")
    sql.append("-- ============================================")
    sql.append("-- DB has Q1378 (ผู้บริหารส่งเสริมการทำงานร่วมกันในระดับชุมชน...) not in Word")
    sql.append("DELETE FROM questions WHERE id = 1378;")
    sql.append("")

    db36_after = [q for q in db36 if q['id'] != 1378]
    for i, wq in enumerate(word_ext912):
        if i < len(db36_after):
            w_text = normalize(wq['text'])
            d_text = normalize(db36_after[i]['text'])
            if w_text != d_text:
                sql.append(f"UPDATE questions SET title = '{escape_sql(w_text)}' WHERE id = {db36_after[i]['id']};")
    sql.append("")

    # ===========================
    # EVAL 37: Internal 4-8 (fix Q10)
    # ===========================
    db37 = get_db_part1_rating_questions(get_db_eval(db_evals, 37))
    sql.append("-- ============================================")
    sql.append("-- EVAL 37: Internal 4-8 (fix Q10)")
    sql.append("-- ============================================")

    for i, wq in enumerate(word_int48):
        if i < len(db37):
            w_text = normalize(wq['text'])
            d_text = normalize(db37[i]['text'])
            if w_text != d_text:
                sql.append(f"UPDATE questions SET title = '{escape_sql(w_text)}' WHERE id = {db37[i]['id']};")
    sql.append("")

    # ===========================
    # EVAL 38: Self Governor (ALL 22 Qs replaced)
    # ===========================
    db38 = get_db_part1_rating_questions(get_db_eval(db_evals, 38))
    sql.append("-- ============================================")
    sql.append("-- EVAL 38: Self Governor (ALL 22 Qs replaced)")
    sql.append("-- ============================================")

    for i, wq in enumerate(word_gov):
        if i < len(db38):
            w_text = normalize(wq['text'])
            sql.append(f"UPDATE questions SET title = '{escape_sql(w_text)}' WHERE id = {db38[i]['id']};")
    sql.append("")

    # ===========================
    # EVAL 39: Self 9-12 (23 in DB -> 30 in Word)
    # ===========================
    eval39 = get_db_eval(db_evals, 39)
    db39 = get_db_part1_rating_questions(eval39)
    part1_39 = [p for p in eval39['parts'] if p['order'] == 1][0]
    part_id_39 = part1_39['id']

    sql.append("-- ============================================")
    sql.append("-- EVAL 39: Self 9-12 (23->30, update + add 7)")
    sql.append("-- ============================================")

    # Update existing 23 questions with correct Word text (same as eval 35)
    for i in range(len(db39)):
        if i < len(word_int912):
            w_text = normalize(word_int912[i]['text'])
            d_text = normalize(db39[i]['text'])
            if w_text != d_text:
                sql.append(f"UPDATE questions SET title = '{escape_sql(w_text)}' WHERE id = {db39[i]['id']};")

    sql.append("")
    sql.append("-- Add missing Interpersonal questions (Q24-26)")

    # Find interpersonal aspect in eval 39
    inter_aspect_39 = None
    for a in part1_39['aspects']:
        if 'ระหว่างบุคคล' in a['name'] or 'Interpersonal' in a['name']:
            inter_aspect_39 = a
            break

    if inter_aspect_39:
        for idx, qi in enumerate([23, 24, 25]):  # word index
            if qi < len(word_int912):
                w_text = normalize(word_int912[qi]['text'])
                sql.append(f"INSERT INTO questions (part_id, aspect_id, sub_aspect_id, title, type, `order`, created_at, updated_at) "
                           f"VALUES ({part_id_39}, {inter_aspect_39['id']}, NULL, '{escape_sql(w_text)}', 'rating', {idx+5}, NOW(), NOW());")

    sql.append("")
    sql.append("-- Add Risk Management aspect + 4 questions (Q27-30)")
    sql.append(f"INSERT INTO aspects (part_id, name, has_subaspects, created_at, updated_at) "
               f"VALUES ({part_id_39}, 'ด้านการบริหารความเสี่ยงและความรับผิดชอบ (Risk Management and Accountability)', 0, NOW(), NOW());")
    sql.append("SET @risk_aspect_39 = LAST_INSERT_ID();")

    for idx, qi in enumerate([26, 27, 28, 29]):
        if qi < len(word_int912):
            w_text = normalize(word_int912[qi]['text'])
            sql.append(f"INSERT INTO questions (part_id, aspect_id, sub_aspect_id, title, type, `order`, created_at, updated_at) "
                       f"VALUES ({part_id_39}, @risk_aspect_39, NULL, '{escape_sql(w_text)}', 'rating', {idx+1}, NOW(), NOW());")

    sql.append("")

    # ===========================
    # EVAL 40: Self 4-8 (13 in DB -> 23 in Word)
    # ===========================
    eval40 = get_db_eval(db_evals, 40)
    db40 = get_db_part1_rating_questions(eval40)
    part1_40 = [p for p in eval40['parts'] if p['order'] == 1][0]
    part_id_40 = part1_40['id']

    sql.append("-- ============================================")
    sql.append("-- EVAL 40: Self 4-8 (13->23, update + add 10)")
    sql.append("-- ============================================")

    # Update existing 13 questions
    for i in range(len(db40)):
        if i < len(word_int48):
            w_text = normalize(word_int48[i]['text'])
            d_text = normalize(db40[i]['text'])
            if w_text != d_text:
                sql.append(f"UPDATE questions SET title = '{escape_sql(w_text)}' WHERE id = {db40[i]['id']};")

    sql.append("")

    # Find aspects for eval 40
    iq_40 = eq_40 = sustainability_40 = None
    for a in part1_40['aspects']:
        if 'เก่งคิด' in a['name'] or 'IQ' in a['name']:
            iq_40 = a
        elif 'เก่งคน' in a['name'] or 'EQ' in a['name']:
            eq_40 = a
        elif 'ยั่งยืน' in a['name'] or 'Sustainability' in a['name']:
            sustainability_40 = a

    # Add missing Q4 (IQ aspect) - "สามารถตรวจสอบความถูกต้อง..."
    sql.append("-- Add missing Q4 for IQ aspect")
    if iq_40:
        w_text = normalize(word_int48[3]['text'])
        sql.append(f"INSERT INTO questions (part_id, aspect_id, sub_aspect_id, title, type, `order`, created_at, updated_at) "
                   f"VALUES ({part_id_40}, {iq_40['id']}, NULL, '{escape_sql(w_text)}', 'rating', 4, NOW(), NOW());")

    # Add missing Q8 (EQ aspect) - "สามารถประสานงานและแลกเปลี่ยนข้อมูล..."
    sql.append("-- Add missing Q8 for EQ aspect")
    if eq_40:
        w_text = normalize(word_int48[7]['text'])
        sql.append(f"INSERT INTO questions (part_id, aspect_id, sub_aspect_id, title, type, `order`, created_at, updated_at) "
                   f"VALUES ({part_id_40}, {eq_40['id']}, NULL, '{escape_sql(w_text)}', 'rating', 4, NOW(), NOW());")

    sql.append("")
    sql.append("-- Add Innovative Thinking aspect + 4 questions (Q16-19)")
    sql.append(f"INSERT INTO aspects (part_id, name, has_subaspects, created_at, updated_at) "
               f"VALUES ({part_id_40}, 'ด้านการคิดเชิงนวัตกรรมและการแก้ปัญหา (Innovative Thinking and Problem Solving)', 0, NOW(), NOW());")
    sql.append("SET @innovative_40 = LAST_INSERT_ID();")

    for idx in range(4):
        qi = 15 + idx  # word index 15-18 = Q16-19
        if qi < len(word_int48):
            w_text = normalize(word_int48[qi]['text'])
            sql.append(f"INSERT INTO questions (part_id, aspect_id, sub_aspect_id, title, type, `order`, created_at, updated_at) "
                       f"VALUES ({part_id_40}, @innovative_40, NULL, '{escape_sql(w_text)}', 'rating', {idx+1}, NOW(), NOW());")

    sql.append("")
    sql.append("-- Add Learning Agility aspect + 4 questions (Q20-23)")
    sql.append(f"INSERT INTO aspects (part_id, name, has_subaspects, created_at, updated_at) "
               f"VALUES ({part_id_40}, 'ด้านการเรียนรู้และการปรับตัวอย่างต่อเนื่อง (Learning Agility and Adaptability)', 0, NOW(), NOW());")
    sql.append("SET @learning_40 = LAST_INSERT_ID();")

    for idx in range(4):
        qi = 19 + idx  # word index 19-22 = Q20-23
        if qi < len(word_int48):
            w_text = normalize(word_int48[qi]['text'])
            sql.append(f"INSERT INTO questions (part_id, aspect_id, sub_aspect_id, title, type, `order`, created_at, updated_at) "
                       f"VALUES ({part_id_40}, @learning_40, NULL, '{escape_sql(w_text)}', 'rating', {idx+1}, NOW(), NOW());")

    sql.append("")
    sql.append("-- ============================================")
    sql.append("-- CLEANUP: Delete year 69 assignments (can re-seed later)")
    sql.append("-- ============================================")
    sql.append("DELETE FROM evaluation_assignments WHERE fiscal_year = 2026;")
    sql.append("")
    sql.append("-- DONE!")

    # Write output
    outpath = os.path.join(os.path.dirname(__file__), 'update_eval_questions.sql')
    with open(outpath, 'w', encoding='utf-8') as f:
        f.write('\n'.join(sql))

    update_count = sum(1 for l in sql if l.startswith('UPDATE'))
    insert_count = sum(1 for l in sql if l.startswith('INSERT'))
    delete_count = sum(1 for l in sql if l.startswith('DELETE'))
    print(f"\nGenerated SQL: {outpath}")
    print(f"  UPDATEs: {update_count}")
    print(f"  INSERTs: {insert_count}")
    print(f"  DELETEs: {delete_count}")
    print(f"  Total lines: {len(sql)}")


if __name__ == '__main__':
    main()
